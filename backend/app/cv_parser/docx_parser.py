import io

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from app.cv_parser.models import CVParseResult
from app.cv_parser.sections import detect_sections


def _has_multi_column(document: DocxDocument) -> bool:
    for section in document.sections:
        for cols in section._sectPr.findall(qn("w:cols")):
            num = cols.get(qn("w:num"))
            if num is not None and int(num) > 1:
                return True
    return False


def _iter_paragraphs(parent: DocxDocument | _Cell) -> list[str]:
    # document.paragraphs only walks top-level body paragraphs, so any CV
    # whose content sits inside a table - a very common way Word templates
    # lay out a sidebar/two-column resume - would extract as empty text and
    # get wrongly rejected as "looks like a scanned image". Walk the body in
    # document order instead, descending into table cells (and their own
    # nested tables) so every paragraph's text is captured regardless of
    # layout.
    body = parent.element.body if isinstance(parent, DocxDocument) else parent._tc
    texts: list[str] = []
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            texts.append(Paragraph(child, parent).text)
        elif isinstance(child, CT_Tbl):
            table = Table(child, parent)
            for row in table.rows:
                for cell in row.cells:
                    texts.extend(_iter_paragraphs(cell))
    return texts


def parse_docx(file_bytes: bytes) -> CVParseResult:
    document = Document(io.BytesIO(file_bytes))
    text = "\n".join(_iter_paragraphs(document))
    return CVParseResult(
        text=text,
        has_tables=len(document.tables) > 0,
        has_multi_column=_has_multi_column(document),
        has_images=len(document.inline_shapes) > 0,
        detected_sections=detect_sections(text),
    )
