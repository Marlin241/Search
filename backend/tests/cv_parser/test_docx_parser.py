import io

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image

from app.cv_parser.docx_parser import parse_docx


def _save(document: Document) -> bytes:
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_parses_simple_docx_and_detects_sections():
    document = Document()
    document.add_paragraph("Expérience professionnelle")
    document.add_paragraph("Développeur chez Acme, 2020-2024")
    document.add_paragraph("Formation")
    document.add_paragraph("Master informatique")
    document.add_paragraph("Compétences")
    document.add_paragraph("Python, FastAPI")

    result = parse_docx(_save(document))

    assert "Développeur chez Acme" in result.text
    assert result.detected_sections == {"experience", "education", "skills"}
    assert result.has_tables is False
    assert result.has_images is False
    assert result.has_multi_column is False


def test_detects_table():
    document = Document()
    document.add_paragraph("Expérience professionnelle")
    document.add_table(rows=2, cols=2)

    result = parse_docx(_save(document))
    assert result.has_tables is True


def test_extracts_text_from_table_only_layout():
    # Many real-world CV templates (two-column "sidebar" layouts in
    # particular) put all of their content inside a table instead of body
    # paragraphs, since Word tables are the common way to lay out columns.
    # document.paragraphs only sees top-level body paragraphs, so a CV like
    # this used to extract as empty text and get rejected as "looks like a
    # scanned image" even though it's fully text-based.
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Jean Dupont\nExpérience professionnelle"
    table.cell(0, 1).text = "Développeur chez Acme, 2020-2024\nFormation\nMaster informatique"

    result = parse_docx(_save(document))

    assert "Jean Dupont" in result.text
    assert "Développeur chez Acme" in result.text
    assert result.detected_sections == {"experience", "education"}


def test_detects_image():
    document = Document()
    document.add_paragraph("Expérience professionnelle")
    img_buffer = io.BytesIO()
    Image.new("RGB", (10, 10), color="black").save(img_buffer, format="PNG")
    img_buffer.seek(0)
    document.add_picture(img_buffer)

    result = parse_docx(_save(document))
    assert result.has_images is True


def test_detects_multi_column_section():
    document = Document()
    document.add_paragraph("Expérience professionnelle")
    section = document.sections[0]
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), "2")
    section._sectPr.append(cols)

    result = parse_docx(_save(document))
    assert result.has_multi_column is True
