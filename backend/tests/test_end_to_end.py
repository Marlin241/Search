import io

from docx import Document

from app.llm_analyzer.analyzer import SemanticReport
from app.llm_analyzer.dependencies import get_semantic_analyzer
from app.main import app


class FakeAnalyzer:
    def analyze(self, cv_text: str, offer_text: str) -> SemanticReport:
        return SemanticReport(score=50, missing_keywords=["Kubernetes"], recommendations=["Learn Kubernetes"])


def _cv_bytes() -> bytes:
    document = Document()
    document.add_paragraph("John Doe").runs[0].bold = True
    document.add_paragraph("john.doe@example.com | +1-555-0100")
    document.add_paragraph("Expérience professionnelle")
    document.add_paragraph("Senior Software Engineer at TechCorp (2019-2024): Led development of microservices architecture using Python and Docker.")
    document.add_paragraph("Software Engineer at StartupXYZ (2017-2019): Developed RESTful APIs using FastAPI and PostgreSQL.")
    document.add_paragraph("Formation")
    document.add_paragraph("Bachelor of Science in Computer Science, University of State (2017)")
    document.add_paragraph("Compétences")
    document.add_paragraph("Python, FastAPI, Docker, Kubernetes, PostgreSQL, AWS, Git, CI/CD")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_cors_allows_configured_origin(client):
    response = client.get("/health", headers={"Origin": "http://localhost:3000"})
    assert response.headers.get("access-control-allow-origin") == "http://localhost:3000"


def test_full_flow_register_login_diagnose_list_delete(client):
    app.dependency_overrides[get_semantic_analyzer] = lambda: FakeAnalyzer()

    register = client.post("/auth/register", json={"email": "flow@example.com", "password": "s3cret!1"})
    assert register.status_code == 201

    login = client.post("/auth/login", data={"username": "flow@example.com", "password": "s3cret!1"})
    assert login.status_code == 200
    headers = {"Authorization": f"Bearer {login.json()['access_token']}"}

    diagnose = client.post(
        "/diagnostics",
        headers=headers,
        files={"cv_file": ("cv.docx", _cv_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        data={"offer_text": "Looking for a Kubernetes engineer."},
    )
    assert diagnose.status_code == 201
    assert diagnose.json()["overall_score"] == 75

    listed = client.get("/diagnostics", headers=headers)
    assert len(listed.json()) == 1

    deleted = client.delete("/diagnostics", headers=headers)
    assert deleted.status_code == 204

    assert client.get("/diagnostics", headers=headers).json() == []

    app.dependency_overrides.pop(get_semantic_analyzer, None)
