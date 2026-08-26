import io

import httpx
import respx
from docx import Document

from app.ats_adapters.custom_fields import CustomFieldAnswerer
from app.ats_adapters.dependencies import get_custom_field_answerer
from app.llm_analyzer.analyzer import SemanticReport
from app.llm_analyzer.dependencies import get_semantic_analyzer
from app.main import app
from app.models.prefilled_form_request_log import PrefilledFormRequestLog
from app.models.user import User
from app.personalization.dependencies import get_cover_letter_generator, get_cv_rewriter
from app.personalization.schemas import CoverLetter, CvExperienceEntry, RewrittenCv
from app.rate_limit.limiter import MAX_PREFILLED_FORM_PREVIEWS_PER_HOUR
from app.storage.client import ObjectStorage, ObjectStorageError
from app.storage.dependencies import get_object_storage

_GREENHOUSE_FORM_HTML = """
<html><body>
<form action="https://boards-api.greenhouse.io/v1/boards/acme/jobs/123" method="post">
  <input type="hidden" name="authenticity_token" value="tok-gh" />
  <label for="first_name">First Name</label>
  <input type="text" name="job_application[first_name]" id="first_name" required />
  <label for="email">Email</label>
  <input type="email" name="job_application[email]" id="email" required />
  <input type="file" name="job_application[resume]" />
  <label for="q1">Why do you want to work here?</label>
  <textarea name="job_application[answers_attributes][0][text_value]" id="q1"></textarea>
</form>
</body></html>
"""


class FakeAnalyzer:
    def analyze(self, cv_text, offer_text):
        return SemanticReport(
            score=70, missing_keywords=["Docker"], recommendations=["Add Docker"]
        )


class FakeCvRewriter:
    def rewrite(self, cv_text, offer_text, missing_keywords, recommendations, **kwargs):
        return RewrittenCv(
            summary="Résumé.",
            experience=[
                CvExperienceEntry(
                    title="Dev",
                    company="Acme",
                    dates="2020-2022",
                    bullets=["A conçu des API."],
                )
            ],
            education=["Master"],
            skills=["Python"],
        )


class FakeHallucinatingCvRewriter:
    """Returns a rewritten CV mentioning an employer and dates that are
    absent from the reference CV, so `cv_needs_review` flags it - the
    anti-hallucination guard from sous-projet 3."""

    def rewrite(self, cv_text, offer_text, missing_keywords, recommendations, **kwargs):
        return RewrittenCv(
            summary="Résumé.",
            experience=[
                CvExperienceEntry(
                    title="Dev",
                    company="Globex Corporation",
                    dates="1998-1999",
                    bullets=["A dirigé Globex Corporation."],
                )
            ],
            education=["Master"],
            skills=["Python"],
        )


class FakeCoverLetterGenerator:
    def generate(
        self, cv_text, offer_text, missing_keywords, recommendations, tone="sobre"
    ):
        return CoverLetter(
            greeting="Madame, Monsieur,",
            body_paragraphs=["Je candidate à ce poste."],
            closing_formula="Cordialement,",
            signature="Jane Doe",
        )


class FakeCustomFieldAnswerer(CustomFieldAnswerer):
    def __init__(self):
        pass

    def answer(self, custom_fields, cv_text, offer_text):
        return {f.name: "Réponse générée." for f in custom_fields}


class FakeObjectStorage(ObjectStorage):
    def __init__(self):
        self._objects: dict[str, bytes] = {}

    def upload(self, key, content):
        self._objects[key] = content

    def download(self, key):
        if key not in self._objects:
            raise ObjectStorageError(f"missing {key}")
        return self._objects[key]

    def delete(self, key):
        self._objects.pop(key, None)


def _register_and_login(client, email: str = "jane@example.com") -> str:
    client.post("/auth/register", json={"email": email, "password": "s3cret!1"})
    login = client.post("/auth/login", data={"username": email, "password": "s3cret!1"})
    return login.json()["access_token"]


def _cv_docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("Jane Doe - Développeuse Python")
    document.add_paragraph("Expérience professionnelle")
    document.add_paragraph(
        "Développeuse Python chez Acme, 2020-2022. A conçu des API REST avec FastAPI."
    )
    document.add_paragraph("Compétences: Python, SQL, Docker.")
    document.add_paragraph("Formation: Master informatique.")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _setup_profile(client, headers: dict) -> None:
    client.put(
        "/profile",
        headers=headers,
        json={
            "first_name": "Jane",
            "last_name": "Doe",
            "phone": "0612345678",
            "work_authorization": "FR/UE",
        },
    )
    client.post(
        "/profile/cv",
        headers=headers,
        files={
            "cv_file": (
                "cv.docx",
                _cv_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )


def _setup_ready_ats_application(client, headers: dict) -> int:
    _setup_profile(client, headers)
    created = client.post(
        "/applications",
        headers=headers,
        json={
            "offer_url": "https://boards.greenhouse.io/acme/jobs/123",
            "offer_text": "Nous recherchons un développeur Python.",
            "source": "greenhouse",
            "company_name": "Acme",
            "job_title": "Développeur Python",
            "ats_type": "greenhouse",
        },
    )
    application_id = created.json()["id"]
    diagnostic_id = created.json()["diagnostic_id"]
    client.post(f"/diagnostics/{diagnostic_id}/cv", headers=headers)
    client.post(f"/diagnostics/{diagnostic_id}/lettre", headers=headers)
    return application_id


def _override_common_dependencies() -> None:
    # get_object_storage is a single @lru_cache singleton for the app's
    # lifetime in production, so uploads made by one request (e.g. CV/lettre
    # generation) are visible to a later request's download (e.g. confirm).
    # A fresh FakeObjectStorage() per Depends() resolution would break that:
    # each request would see an empty store. One shared instance per test
    # reproduces the real singleton behavior.
    fake_storage = FakeObjectStorage()
    app.dependency_overrides[get_semantic_analyzer] = lambda: FakeAnalyzer()
    app.dependency_overrides[get_cv_rewriter] = lambda: FakeCvRewriter()
    app.dependency_overrides[get_cover_letter_generator] = lambda: (
        FakeCoverLetterGenerator()
    )
    app.dependency_overrides[get_object_storage] = lambda: fake_storage
    app.dependency_overrides[get_custom_field_answerer] = lambda: (
        FakeCustomFieldAnswerer()
    )


@respx.mock
def test_get_prefilled_form_returns_standard_and_llm_answered_custom_fields(client):
    _override_common_dependencies()
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    application_id = _setup_ready_ats_application(client, headers)
    respx.get("https://boards.greenhouse.io/acme/jobs/123").mock(
        return_value=httpx.Response(200, text=_GREENHOUSE_FORM_HTML)
    )

    response = client.get(
        f"/applications/{application_id}/prefilled-form", headers=headers
    )

    assert response.status_code == 200
    fields = response.json()["fields"]
    first_name = next(f for f in fields if f["name"] == "job_application[first_name]")
    assert first_name["value"] == "Jane"
    custom = next(f for f in fields if f["is_custom"])
    assert custom["value"] == "Réponse générée."


def test_get_prefilled_form_returns_409_for_non_ats_offer(client):
    _override_common_dependencies()
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    _setup_profile(client, headers)
    created = client.post(
        "/applications",
        headers=headers,
        json={
            "offer_url": "https://www.linkedin.com/jobs/view/123",
            "offer_text": "Offre.",
            "source": "manual",
            "company_name": "Acme",
            "job_title": "Dev",
        },
    )
    response = client.get(
        f"/applications/{created.json()['id']}/prefilled-form", headers=headers
    )
    assert response.status_code == 409


@respx.mock
def test_confirm_application_auto_submits_for_ats_offer(client):
    _override_common_dependencies()
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    application_id = _setup_ready_ats_application(client, headers)
    respx.get("https://boards.greenhouse.io/acme/jobs/123").mock(
        return_value=httpx.Response(200, text=_GREENHOUSE_FORM_HTML)
    )
    submit_route = respx.post(
        "https://boards-api.greenhouse.io/v1/boards/acme/jobs/123"
    ).mock(return_value=httpx.Response(200))

    prefilled = client.get(
        f"/applications/{application_id}/prefilled-form", headers=headers
    ).json()
    response = client.post(
        f"/applications/{application_id}/confirm",
        headers=headers,
        json={"fields": prefilled["fields"]},
    )

    assert response.status_code == 200
    assert response.json()["status"] == "soumise_auto"
    assert submit_route.called


@respx.mock
def test_confirm_application_records_failure_status_on_submission_error(
    client, db_session
):
    _override_common_dependencies()
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    application_id = _setup_ready_ats_application(client, headers)
    respx.get("https://boards.greenhouse.io/acme/jobs/123").mock(
        return_value=httpx.Response(200, text=_GREENHOUSE_FORM_HTML)
    )
    respx.post("https://boards-api.greenhouse.io/v1/boards/acme/jobs/123").mock(
        return_value=httpx.Response(500)
    )

    prefilled = client.get(
        f"/applications/{application_id}/prefilled-form", headers=headers
    ).json()
    response = client.post(
        f"/applications/{application_id}/confirm",
        headers=headers,
        json={"fields": prefilled["fields"]},
    )

    assert response.status_code == 503
    # The `client` fixture shares one db_session across every request in
    # this test, so without this, the follow-up GET below would just
    # re-read the same in-memory identity-mapped object `confirm` already
    # mutated in Python - a completely inert check for whether the failure
    # status was actually committed to the database. expire_all() forces
    # the next access to hit the DB for real.
    db_session.expire_all()
    detail = client.get(f"/applications/{application_id}", headers=headers).json()
    assert detail["status"] == "echec_soumission"
    assert detail["error_message"] is not None


def test_confirm_application_without_ats_type_moves_to_assisted_status(client):
    _override_common_dependencies()
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    _setup_profile(client, headers)
    created = client.post(
        "/applications",
        headers=headers,
        json={
            "offer_url": "https://www.linkedin.com/jobs/view/123",
            "offer_text": "Offre.",
            "source": "manual",
            "company_name": "Acme",
            "job_title": "Dev",
        },
    )
    application_id = created.json()["id"]

    response = client.post(
        f"/applications/{application_id}/confirm", headers=headers, json={}
    )

    assert response.status_code == 200
    assert response.json()["status"] == "a_soumettre_manuellement"


def test_mark_sent_manually_transitions_status(client):
    _override_common_dependencies()
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    _setup_profile(client, headers)
    created = client.post(
        "/applications",
        headers=headers,
        json={
            "offer_url": "https://www.linkedin.com/jobs/view/123",
            "offer_text": "Offre.",
            "source": "manual",
            "company_name": "Acme",
            "job_title": "Dev",
        },
    )
    application_id = created.json()["id"]
    client.post(f"/applications/{application_id}/confirm", headers=headers, json={})

    response = client.post(f"/applications/{application_id}/mark-sent", headers=headers)

    assert response.status_code == 200
    assert response.json()["status"] == "soumise_manuelle_confirmee"


def test_mark_sent_manually_rejects_wrong_state(client):
    _override_common_dependencies()
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    application_id = _setup_ready_ats_application(
        client, headers
    )  # still "en_cours", never confirmed

    response = client.post(f"/applications/{application_id}/mark-sent", headers=headers)
    assert response.status_code == 409


def test_prefilled_form_and_confirm_handle_unknown_ats_type_gracefully(client):
    # ats_type is an unvalidated free-form string on ApplicationCreateIn - an
    # application can end up with an ats_type that isn't in the adapter
    # registry (e.g. "workday", which has no adapter implemented). Both
    # endpoints must treat "no adapter available" the same as "no ats_type
    # at all" rather than crashing on a None adapter.
    _override_common_dependencies()
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    _setup_profile(client, headers)
    created = client.post(
        "/applications",
        headers=headers,
        json={
            "offer_url": "https://workday.example.com/jobs/123",
            "offer_text": "Offre.",
            "source": "workday",
            "company_name": "Acme",
            "job_title": "Dev",
            "ats_type": "workday",
        },
    )
    application_id = created.json()["id"]

    prefilled_response = client.get(
        f"/applications/{application_id}/prefilled-form", headers=headers
    )
    assert prefilled_response.status_code == 409

    confirm_response = client.post(
        f"/applications/{application_id}/confirm", headers=headers, json={}
    )
    assert confirm_response.status_code == 200
    assert confirm_response.json()["status"] == "a_soumettre_manuellement"


@respx.mock
def test_confirm_application_second_attempt_after_success_is_rejected(client):
    # Proves the status-check-then-transition ordering: once the first
    # confirm has succeeded (status is no longer en_cours), a second confirm
    # attempt on the same application must be rejected rather than
    # resubmitting to the employer's ATS. This can't exercise genuine
    # multi-threaded concurrency against SQLite, but it does prove that the
    # row-lock-then-refresh in confirm_application picks up the committed
    # status change rather than reusing a stale read.
    _override_common_dependencies()
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    application_id = _setup_ready_ats_application(client, headers)
    respx.get("https://boards.greenhouse.io/acme/jobs/123").mock(
        return_value=httpx.Response(200, text=_GREENHOUSE_FORM_HTML)
    )
    submit_route = respx.post(
        "https://boards-api.greenhouse.io/v1/boards/acme/jobs/123"
    ).mock(return_value=httpx.Response(200))

    prefilled = client.get(
        f"/applications/{application_id}/prefilled-form", headers=headers
    ).json()
    first = client.post(
        f"/applications/{application_id}/confirm",
        headers=headers,
        json={"fields": prefilled["fields"]},
    )
    assert first.status_code == 200
    assert first.json()["status"] == "soumise_auto"
    assert submit_route.call_count == 1

    second = client.post(
        f"/applications/{application_id}/confirm",
        headers=headers,
        json={"fields": prefilled["fields"]},
    )

    assert second.status_code == 409
    assert submit_route.call_count == 1


@respx.mock
def test_confirm_application_can_be_retried_after_a_failed_submission(client):
    # Finding 2: `echec_soumission` used to be a permanent dead end - confirm
    # required `en_cours`, mark-sent required `a_soumettre_manuellement`, and
    # the (user_id, offer_url) unique constraint blocked re-creating. A
    # transient network blip during submit therefore stranded the candidature
    # forever. Retrying is user-initiated, so it does not violate the
    # "no automatic retry" constraint.
    _override_common_dependencies()
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    application_id = _setup_ready_ats_application(client, headers)
    respx.get("https://boards.greenhouse.io/acme/jobs/123").mock(
        return_value=httpx.Response(200, text=_GREENHOUSE_FORM_HTML)
    )
    submit_route = respx.post(
        "https://boards-api.greenhouse.io/v1/boards/acme/jobs/123"
    ).mock(side_effect=[httpx.Response(500), httpx.Response(200)])

    prefilled = client.get(
        f"/applications/{application_id}/prefilled-form", headers=headers
    ).json()

    first = client.post(
        f"/applications/{application_id}/confirm",
        headers=headers,
        json={"fields": prefilled["fields"]},
    )
    assert first.status_code == 503
    assert (
        client.get(f"/applications/{application_id}", headers=headers).json()["status"]
        == "echec_soumission"
    )

    second = client.post(
        f"/applications/{application_id}/confirm",
        headers=headers,
        json={"fields": prefilled["fields"]},
    )

    assert second.status_code == 200
    assert second.json()["status"] == "soumise_auto"
    assert second.json()["error_message"] is None
    assert submit_route.call_count == 2


def test_confirm_application_rejects_retry_from_terminal_statuses(client):
    # The retry allowance of Finding 2 is narrow: only `echec_soumission`
    # joins `en_cours` as a valid starting state. `a_soumettre_manuellement`
    # must still be rejected (it is handled by mark-sent, not confirm).
    _override_common_dependencies()
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    _setup_profile(client, headers)
    created = client.post(
        "/applications",
        headers=headers,
        json={
            "offer_url": "https://www.linkedin.com/jobs/view/123",
            "offer_text": "Offre.",
            "source": "manual",
            "company_name": "Acme",
            "job_title": "Dev",
        },
    )
    application_id = created.json()["id"]
    assert (
        client.post(
            f"/applications/{application_id}/confirm", headers=headers, json={}
        ).status_code
        == 200
    )

    response = client.post(
        f"/applications/{application_id}/confirm", headers=headers, json={}
    )

    assert response.status_code == 409


@respx.mock
def test_confirm_application_blocks_auto_submit_when_cv_needs_review(client):
    # Finding 5: `needs_review` is the anti-hallucination flag - a CV that
    # mentions employers/schools/dates absent from the reference CV must
    # never be auto-submitted to a real employer unreviewed.
    _override_common_dependencies()
    app.dependency_overrides[get_cv_rewriter] = lambda: FakeHallucinatingCvRewriter()
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    application_id = _setup_ready_ats_application(client, headers)
    respx.get("https://boards.greenhouse.io/acme/jobs/123").mock(
        return_value=httpx.Response(200, text=_GREENHOUSE_FORM_HTML)
    )
    submit_route = respx.post(
        "https://boards-api.greenhouse.io/v1/boards/acme/jobs/123"
    ).mock(return_value=httpx.Response(200))

    prefilled = client.get(
        f"/applications/{application_id}/prefilled-form", headers=headers
    ).json()
    response = client.post(
        f"/applications/{application_id}/confirm",
        headers=headers,
        json={"fields": prefilled["fields"]},
    )

    assert response.status_code == 422
    assert "vérifier" in response.json()["detail"]
    assert not submit_route.called
    assert (
        client.get(f"/applications/{application_id}", headers=headers).json()["status"]
        == "en_cours"
    )


@respx.mock
def test_confirm_application_proceeds_when_cv_does_not_need_review(client):
    _override_common_dependencies()  # the default FakeCvRewriter stays faithful to the reference CV
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    application_id = _setup_ready_ats_application(client, headers)
    respx.get("https://boards.greenhouse.io/acme/jobs/123").mock(
        return_value=httpx.Response(200, text=_GREENHOUSE_FORM_HTML)
    )
    submit_route = respx.post(
        "https://boards-api.greenhouse.io/v1/boards/acme/jobs/123"
    ).mock(return_value=httpx.Response(200))

    prefilled = client.get(
        f"/applications/{application_id}/prefilled-form", headers=headers
    ).json()
    response = client.post(
        f"/applications/{application_id}/confirm",
        headers=headers,
        json={"fields": prefilled["fields"]},
    )

    assert response.status_code == 200
    assert response.json()["status"] == "soumise_auto"
    assert submit_route.called


@respx.mock
def test_confirm_application_override_needs_review_allows_auto_submit(client):
    # The user has manually read the flagged CV and judged it fine: passing
    # override_needs_review=True must let the auto-submit path through even
    # though needs_review is still True on the document.
    _override_common_dependencies()
    app.dependency_overrides[get_cv_rewriter] = lambda: FakeHallucinatingCvRewriter()
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    application_id = _setup_ready_ats_application(client, headers)
    respx.get("https://boards.greenhouse.io/acme/jobs/123").mock(
        return_value=httpx.Response(200, text=_GREENHOUSE_FORM_HTML)
    )
    submit_route = respx.post(
        "https://boards-api.greenhouse.io/v1/boards/acme/jobs/123"
    ).mock(return_value=httpx.Response(200))

    prefilled = client.get(
        f"/applications/{application_id}/prefilled-form", headers=headers
    ).json()
    response = client.post(
        f"/applications/{application_id}/confirm",
        headers=headers,
        json={"fields": prefilled["fields"], "override_needs_review": True},
    )

    assert response.status_code == 200
    assert response.json()["status"] == "soumise_auto"
    assert submit_route.called


@respx.mock
def test_confirm_application_override_needs_review_is_noop_when_not_flagged(client):
    # override_needs_review=True with a CV that doesn't need review must not
    # change behavior - it's an override of the block, not a bypass of
    # anything else.
    _override_common_dependencies()  # the default FakeCvRewriter stays faithful to the reference CV
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    application_id = _setup_ready_ats_application(client, headers)
    respx.get("https://boards.greenhouse.io/acme/jobs/123").mock(
        return_value=httpx.Response(200, text=_GREENHOUSE_FORM_HTML)
    )
    submit_route = respx.post(
        "https://boards-api.greenhouse.io/v1/boards/acme/jobs/123"
    ).mock(return_value=httpx.Response(200))

    prefilled = client.get(
        f"/applications/{application_id}/prefilled-form", headers=headers
    ).json()
    response = client.post(
        f"/applications/{application_id}/confirm",
        headers=headers,
        json={"fields": prefilled["fields"], "override_needs_review": True},
    )

    assert response.status_code == 200
    assert response.json()["status"] == "soumise_auto"
    assert submit_route.called


@respx.mock
def test_confirm_application_override_needs_review_does_not_bypass_other_guards(client):
    # override_needs_review must lift *only* the needs_review block. Proving
    # this requires tripping a *different* guard while the override is set,
    # and confirming that guard still fires - a happy-path override run
    # (as in the two tests above) can't distinguish "scoped override" from
    # "override disables all checks", since no other guard is triggered.
    _override_common_dependencies()
    app.dependency_overrides[get_cv_rewriter] = lambda: FakeHallucinatingCvRewriter()
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    application_id = _setup_ready_ats_application(client, headers)
    respx.get("https://boards.greenhouse.io/acme/jobs/123").mock(
        return_value=httpx.Response(200, text=_GREENHOUSE_FORM_HTML)
    )
    submit_route = respx.post(
        "https://boards-api.greenhouse.io/v1/boards/acme/jobs/123"
    ).mock(return_value=httpx.Response(200))

    client.get(f"/applications/{application_id}/prefilled-form", headers=headers)
    # Trip the separate "fields are required" guard (payload.fields is None)
    # while also passing override_needs_review=True: the override must not
    # paper over this unrelated 422.
    response = client.post(
        f"/applications/{application_id}/confirm",
        headers=headers,
        json={"override_needs_review": True},
    )

    assert response.status_code == 422
    assert "champs du formulaire" in response.json()["detail"]
    assert not submit_route.called
    assert (
        client.get(f"/applications/{application_id}", headers=headers).json()["status"]
        == "en_cours"
    )


def test_confirm_application_does_not_block_assisted_mode_when_cv_needs_review(client):
    # Assisted mode (no ats_type) never posts to an employer from the
    # backend - the user submits manually after seeing the "needs review"
    # badge - so the needs_review block must not apply there.
    _override_common_dependencies()
    app.dependency_overrides[get_cv_rewriter] = lambda: FakeHallucinatingCvRewriter()
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    _setup_profile(client, headers)
    created = client.post(
        "/applications",
        headers=headers,
        json={
            "offer_url": "https://www.linkedin.com/jobs/view/123",
            "offer_text": "Offre.",
            "source": "manual",
            "company_name": "Acme",
            "job_title": "Dev",
        },
    )
    application_id = created.json()["id"]
    diagnostic_id = created.json()["diagnostic_id"]
    client.post(f"/diagnostics/{diagnostic_id}/cv", headers=headers)
    client.post(f"/diagnostics/{diagnostic_id}/lettre", headers=headers)

    response = client.post(
        f"/applications/{application_id}/confirm", headers=headers, json={}
    )

    assert response.status_code == 200
    assert response.json()["status"] == "a_soumettre_manuellement"


@respx.mock
def test_get_prefilled_form_is_rate_limited(client, db_session):
    # Finding 3: this is an LLM-calling endpoint (CustomFieldAnswerer) and
    # was the only one with no rate-limit counter of its own.
    _override_common_dependencies()
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    application_id = _setup_ready_ats_application(client, headers)
    respx.get("https://boards.greenhouse.io/acme/jobs/123").mock(
        return_value=httpx.Response(200, text=_GREENHOUSE_FORM_HTML)
    )

    user_id = db_session.query(User).filter(User.email == "jane@example.com").first().id
    for _ in range(MAX_PREFILLED_FORM_PREVIEWS_PER_HOUR):
        db_session.add(PrefilledFormRequestLog(user_id=user_id))
    db_session.commit()

    response = client.get(
        f"/applications/{application_id}/prefilled-form", headers=headers
    )

    assert response.status_code == 429


@respx.mock
def test_get_prefilled_form_records_one_request_log_per_call(client, db_session):
    _override_common_dependencies()
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    application_id = _setup_ready_ats_application(client, headers)
    respx.get("https://boards.greenhouse.io/acme/jobs/123").mock(
        return_value=httpx.Response(200, text=_GREENHOUSE_FORM_HTML)
    )

    assert (
        client.get(
            f"/applications/{application_id}/prefilled-form", headers=headers
        ).status_code
        == 200
    )
    assert (
        client.get(
            f"/applications/{application_id}/prefilled-form", headers=headers
        ).status_code
        == 200
    )

    assert db_session.query(PrefilledFormRequestLog).count() == 2
