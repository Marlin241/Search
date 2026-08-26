import io

from docx import Document

from app.cv_parser.parser import MAX_CV_SIZE_BYTES
from app.llm_analyzer.analyzer import SemanticReport
from app.llm_analyzer.dependencies import get_semantic_analyzer
from app.main import app
from app.models.personalized_document import PersonalizedDocument
from app.personalization.dependencies import get_cv_rewriter
from app.personalization.schemas import CvExperienceEntry, RewrittenCv
from app.rate_limit.limiter import MAX_DIAGNOSTICS_PER_HOUR
from app.storage.client import ObjectStorage, ObjectStorageError
from app.storage.dependencies import get_object_storage


class FakeAnalyzer:
    def analyze(self, cv_text: str, offer_text: str) -> SemanticReport:
        return SemanticReport(
            score=60, missing_keywords=["Docker"], recommendations=["Add Docker"]
        )


def _clean_cv_docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("Expérience professionnelle")
    document.add_paragraph("Développeur")
    document.add_paragraph("Formation")
    document.add_paragraph("Master")
    document.add_paragraph("Compétences")
    document.add_paragraph("Python")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _register_and_login(client, email: str = "jane@example.com") -> str:
    client.post("/auth/register", json={"email": email, "password": "s3cret!1"})
    login = client.post("/auth/login", data={"username": email, "password": "s3cret!1"})
    return login.json()["access_token"]


def test_create_diagnostic_returns_combined_report(client):
    app.dependency_overrides[get_semantic_analyzer] = lambda: FakeAnalyzer()
    token = _register_and_login(client)

    response = client.post(
        "/diagnostics",
        headers={"Authorization": f"Bearer {token}"},
        files={
            "cv_file": (
                "cv.docx",
                _clean_cv_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"offer_text": "We need a Python developer with Docker experience."},
    )

    assert response.status_code == 201
    body = response.json()
    assert body["structural_score"] == 100
    assert body["semantic_score"] == 60
    assert body["overall_score"] == 80
    assert body["missing_keywords"] == ["Docker"]
    assert isinstance(body["id"], int)
    assert body["created_at"]

    app.dependency_overrides.pop(get_semantic_analyzer, None)


def test_create_diagnostic_without_offer_returns_422(client):
    app.dependency_overrides[get_semantic_analyzer] = lambda: FakeAnalyzer()
    token = _register_and_login(client)

    response = client.post(
        "/diagnostics",
        headers={"Authorization": f"Bearer {token}"},
        files={
            "cv_file": (
                "cv.docx",
                _clean_cv_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 422
    app.dependency_overrides.pop(get_semantic_analyzer, None)


def test_create_diagnostic_without_cv_file_reuses_reference_cv(client):
    app.dependency_overrides[get_semantic_analyzer] = lambda: FakeAnalyzer()
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}

    client.put(
        "/profile",
        headers=headers,
        json={
            "first_name": "Jane",
            "last_name": "Doe",
            "phone": "0612345678",
            "work_authorization": "FR/UE",
        },
    )
    client.post(
        "/profile/cv",
        headers=headers,
        files={
            "cv_file": (
                "cv.docx",
                _clean_cv_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    response = client.post(
        "/diagnostics",
        headers=headers,
        data={"offer_text": "We need a Python developer with Docker experience."},
    )

    assert response.status_code == 201
    body = response.json()
    assert body["semantic_score"] == 60

    app.dependency_overrides.pop(get_semantic_analyzer, None)


def test_create_diagnostic_without_cv_file_or_reference_cv_returns_422(client):
    app.dependency_overrides[get_semantic_analyzer] = lambda: FakeAnalyzer()
    token = _register_and_login(client)

    response = client.post(
        "/diagnostics",
        headers={"Authorization": f"Bearer {token}"},
        data={"offer_text": "We need a Python developer with Docker experience."},
    )

    assert response.status_code == 422
    app.dependency_overrides.pop(get_semantic_analyzer, None)


def test_list_and_delete_diagnostics(client):
    app.dependency_overrides[get_semantic_analyzer] = lambda: FakeAnalyzer()
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}

    client.post(
        "/diagnostics",
        headers=headers,
        files={
            "cv_file": (
                "cv.docx",
                _clean_cv_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"offer_text": "We need a Python developer."},
    )

    listed = client.get("/diagnostics", headers=headers)
    assert listed.status_code == 200
    assert len(listed.json()) == 1

    deleted = client.delete("/diagnostics", headers=headers)
    assert deleted.status_code == 204

    listed_after = client.get("/diagnostics", headers=headers)
    assert listed_after.json() == []

    app.dependency_overrides.pop(get_semantic_analyzer, None)


def test_create_diagnostic_oversized_cv_returns_422(client):
    app.dependency_overrides[get_semantic_analyzer] = lambda: FakeAnalyzer()
    token = _register_and_login(client)

    oversized_bytes = b"a" * (MAX_CV_SIZE_BYTES + 1)
    response = client.post(
        "/diagnostics",
        headers={"Authorization": f"Bearer {token}"},
        files={
            "cv_file": (
                "cv.docx",
                oversized_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"offer_text": "We need a Python developer with Docker experience."},
    )

    assert response.status_code == 422

    app.dependency_overrides.pop(get_semantic_analyzer, None)


def test_create_diagnostic_oversized_offer_text_returns_422(client):
    app.dependency_overrides[get_semantic_analyzer] = lambda: FakeAnalyzer()
    token = _register_and_login(client)

    response = client.post(
        "/diagnostics",
        headers={"Authorization": f"Bearer {token}"},
        files={
            "cv_file": (
                "cv.docx",
                _clean_cv_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"offer_text": "a" * 50001},
    )

    assert response.status_code == 422

    app.dependency_overrides.pop(get_semantic_analyzer, None)


def test_rate_limit_returns_429(client):
    app.dependency_overrides[get_semantic_analyzer] = lambda: FakeAnalyzer()
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}

    for _ in range(MAX_DIAGNOSTICS_PER_HOUR):
        response = client.post(
            "/diagnostics",
            headers=headers,
            files={
                "cv_file": (
                    "cv.docx",
                    _clean_cv_docx_bytes(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            data={"offer_text": "We need a Python developer."},
        )
        assert response.status_code == 201

    blocked = client.post(
        "/diagnostics",
        headers=headers,
        files={
            "cv_file": (
                "cv.docx",
                _clean_cv_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"offer_text": "We need a Python developer."},
    )
    assert blocked.status_code == 429

    app.dependency_overrides.pop(get_semantic_analyzer, None)


def test_list_diagnostics_includes_id_and_created_at_newest_first(client):
    app.dependency_overrides[get_semantic_analyzer] = lambda: FakeAnalyzer()
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}

    first = client.post(
        "/diagnostics",
        headers=headers,
        files={
            "cv_file": (
                "cv.docx",
                _clean_cv_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"offer_text": "We need a Python developer."},
    ).json()
    second = client.post(
        "/diagnostics",
        headers=headers,
        files={
            "cv_file": (
                "cv.docx",
                _clean_cv_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"offer_text": "We need a Python developer."},
    ).json()

    listed = client.get("/diagnostics", headers=headers).json()

    assert [d["id"] for d in listed] == [second["id"], first["id"]]
    assert all("created_at" in d for d in listed)

    app.dependency_overrides.pop(get_semantic_analyzer, None)


class _FakeCvRewriter:
    def rewrite(self, cv_text, offer_text, missing_keywords, recommendations, **kwargs):
        return RewrittenCv(
            summary="Résumé.",
            experience=[
                CvExperienceEntry(
                    title="Dev", company="Acme", dates="2020-2022", bullets=["Bullet."]
                )
            ],
            education=["Master"],
            skills=["Python"],
        )


class _FakeObjectStorage(ObjectStorage):
    def __init__(self):
        self._objects: dict[str, bytes] = {}

    def upload(self, key: str, content: bytes) -> None:
        self._objects[key] = content

    def download(self, key: str) -> bytes:
        if key not in self._objects:
            raise ObjectStorageError(f"missing key {key}")
        return self._objects[key]

    def delete(self, key: str) -> None:
        self._objects.pop(key, None)


def test_delete_all_diagnostics_also_purges_personalized_documents(client, db_session):
    app.dependency_overrides[get_semantic_analyzer] = lambda: FakeAnalyzer()
    fake_storage = _FakeObjectStorage()
    app.dependency_overrides[get_object_storage] = lambda: fake_storage
    app.dependency_overrides[get_cv_rewriter] = lambda: _FakeCvRewriter()

    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}

    diagnostic_id = client.post(
        "/diagnostics",
        headers=headers,
        files={
            "cv_file": (
                "cv.docx",
                _clean_cv_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"offer_text": "We need a Python developer."},
    ).json()["id"]

    client.post(f"/diagnostics/{diagnostic_id}/cv", headers=headers)
    assert len(fake_storage._objects) == 1

    deleted = client.delete("/diagnostics", headers=headers)
    assert deleted.status_code == 204

    assert db_session.query(PersonalizedDocument).count() == 0
    assert len(fake_storage._objects) == 0

    app.dependency_overrides.pop(get_semantic_analyzer, None)
    app.dependency_overrides.pop(get_object_storage, None)
    app.dependency_overrides.pop(get_cv_rewriter, None)


def test_delete_all_diagnostics_also_purges_applications(client, db_session):
    from app.models.application import APPLICATION_STATUS_EN_COURS, Application
    from app.models.user import User

    app.dependency_overrides[get_semantic_analyzer] = lambda: FakeAnalyzer()
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}

    create_response = client.post(
        "/diagnostics",
        headers=headers,
        files={
            "cv_file": (
                "cv.docx",
                _clean_cv_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"offer_text": "Nous recherchons un développeur Python."},
    )
    diagnostic_id = create_response.json()["id"]

    user = db_session.query(User).filter(User.email == "jane@example.com").first()
    db_session.add(
        Application(
            user_id=user.id,
            diagnostic_id=diagnostic_id,
            offer_url="https://example.com/job/1",
            source="manual",
            company_name="Acme",
            job_title="Dev",
            ats_type=None,
            status=APPLICATION_STATUS_EN_COURS,
        )
    )
    db_session.commit()

    response = client.delete("/diagnostics", headers=headers)

    assert response.status_code == 204
    assert db_session.query(Application).count() == 0


def test_delete_all_diagnostics_does_not_purge_other_users_applications(
    client, db_session
):
    from app.models.application import APPLICATION_STATUS_EN_COURS, Application
    from app.models.user import User

    app.dependency_overrides[get_semantic_analyzer] = lambda: FakeAnalyzer()

    owner_token = _register_and_login(client, "owner@example.com")
    owner_headers = {"Authorization": f"Bearer {owner_token}"}
    owner_diagnostic_id = client.post(
        "/diagnostics",
        headers=owner_headers,
        files={
            "cv_file": (
                "cv.docx",
                _clean_cv_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"offer_text": "Nous recherchons un développeur Python."},
    ).json()["id"]

    other_token = _register_and_login(client, "other@example.com")
    other_headers = {"Authorization": f"Bearer {other_token}"}
    other_diagnostic_id = client.post(
        "/diagnostics",
        headers=other_headers,
        files={
            "cv_file": (
                "cv.docx",
                _clean_cv_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"offer_text": "Nous recherchons un développeur Java."},
    ).json()["id"]

    owner = db_session.query(User).filter(User.email == "owner@example.com").first()
    other = db_session.query(User).filter(User.email == "other@example.com").first()

    db_session.add(
        Application(
            user_id=owner.id,
            diagnostic_id=owner_diagnostic_id,
            offer_url="https://example.com/job/owner",
            source="manual",
            company_name="Acme",
            job_title="Dev",
            ats_type=None,
            status=APPLICATION_STATUS_EN_COURS,
        )
    )
    other_application = Application(
        user_id=other.id,
        diagnostic_id=other_diagnostic_id,
        offer_url="https://example.com/job/other",
        source="manual",
        company_name="Other Co",
        job_title="Backend Dev",
        ats_type=None,
        status=APPLICATION_STATUS_EN_COURS,
    )
    db_session.add(other_application)
    db_session.commit()
    other_application_id = other_application.id

    response = client.delete("/diagnostics", headers=owner_headers)

    assert response.status_code == 204
    # Only the requesting user's own Application row was purged.
    assert (
        db_session.query(Application).filter(Application.user_id == owner.id).count()
        == 0
    )
    # The other user's Application, attached to a different diagnostic,
    # must survive untouched.
    surviving = (
        db_session.query(Application)
        .filter(Application.id == other_application_id)
        .first()
    )
    assert surviving is not None
    assert surviving.user_id == other.id

    # And it's still visible to that other user through the normal API.
    other_get = client.get("/applications", headers=other_headers)
    assert other_get.status_code == 200
    assert len(other_get.json()) == 1
