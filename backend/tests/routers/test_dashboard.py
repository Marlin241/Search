from app.llm_analyzer.analyzer import SemanticReport
from app.llm_analyzer.dependencies import get_semantic_analyzer
from app.main import app

_SAVED_JOB_URL = "ftp://example.com/job/saved"


class FakeAnalyzer:
    def analyze(self, cv_text, offer_text):
        return SemanticReport(
            score=70, missing_keywords=["Docker"], recommendations=["Add Docker"]
        )


def _register_and_login(client, email: str = "jane@example.com") -> str:
    client.post("/auth/register", json={"email": email, "password": "s3cret!1"})
    login = client.post("/auth/login", data={"username": email, "password": "s3cret!1"})
    return login.json()["access_token"]


def _setup_profile(client, headers: dict) -> None:
    client.put(
        "/profile",
        headers=headers,
        json={
            "first_name": "Jane",
            "last_name": "Doe",
            "phone": "0600000000",
            "work_authorization": "FR/UE",
        },
    )
    import io

    from docx import Document

    document = Document()
    document.add_paragraph("Expérience professionnelle")
    document.add_paragraph("Développeuse Full Stack chez Acme, 2020-2022")
    document.add_paragraph("Formation")
    document.add_paragraph("Master Informatique")
    document.add_paragraph("Compétences")
    document.add_paragraph("Python, Docker")
    buffer = io.BytesIO()
    document.save(buffer)
    client.post(
        "/profile/cv",
        headers=headers,
        files={
            "cv_file": (
                "cv.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )


def _create_application(client, headers: dict, offer_url: str) -> dict:
    response = client.post(
        "/applications",
        headers=headers,
        json={
            "offer_url": offer_url,
            "offer_text": "Offre.",
            "source": "manual",
            "company_name": "Acme",
            "job_title": "Dev",
        },
    )
    assert response.status_code == 201
    return response.json()


def test_update_funnel_stage_and_ownership(client):
    app.dependency_overrides[get_semantic_analyzer] = lambda: FakeAnalyzer()
    owner_token = _register_and_login(client, "owner@example.com")
    owner_headers = {"Authorization": f"Bearer {owner_token}"}
    _setup_profile(client, owner_headers)
    application = _create_application(
        client, owner_headers, "https://example.com/job/1"
    )

    response = client.patch(
        f"/applications/{application['id']}/funnel-stage",
        headers=owner_headers,
        json={"funnel_stage": "entretien_programme"},
    )
    assert response.status_code == 200
    assert response.json()["funnel_stage"] == "entretien_programme"

    attacker_token = _register_and_login(client, "attacker@example.com")
    attacker_headers = {"Authorization": f"Bearer {attacker_token}"}
    forbidden = client.patch(
        f"/applications/{application['id']}/funnel-stage",
        headers=attacker_headers,
        json={"funnel_stage": "proposition"},
    )
    assert forbidden.status_code == 404

    app.dependency_overrides.pop(get_semantic_analyzer, None)


def test_dashboard_kanban_groups_by_stage_and_excludes_applied_saved_jobs(client):
    app.dependency_overrides[get_semantic_analyzer] = lambda: FakeAnalyzer()
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    _setup_profile(client, headers)

    applied = _create_application(client, headers, "https://example.com/job/applied")
    client.patch(
        f"/applications/{applied['id']}/funnel-stage",
        headers=headers,
        json={"funnel_stage": "proposition"},
    )

    # A SavedJob whose offer_url matches an existing Application must NOT
    # show up under "sauvegardees" - it already has a real application.
    client.post(
        "/saved-jobs",
        headers=headers,
        json={
            "offer_url": "https://example.com/job/applied",
            "title": "Dev",
            "company": "Acme",
            "location": "Paris",
            "snippet": "Poste.",
            "source": "manual",
        },
    )
    # A SavedJob with no matching Application must show up.
    client.post(
        "/saved-jobs",
        headers=headers,
        json={
            "offer_url": _SAVED_JOB_URL,
            "title": "Développeur Backend",
            "company": "Beta",
            "location": "Lyon",
            "snippet": "Poste.",
            "source": "manual",
        },
    )

    board = client.get("/dashboard/kanban", headers=headers)
    assert board.status_code == 200
    body = board.json()
    assert len(body["proposition"]) == 1
    assert body["proposition"][0]["id"] == applied["id"]
    assert body["postule"] == []
    saved_urls = {sj["offer_url"] for sj in body["sauvegardees"]}
    assert saved_urls == {_SAVED_JOB_URL}

    app.dependency_overrides.pop(get_semantic_analyzer, None)


def test_dashboard_calendar_filters_by_month_and_validates_param(client):
    app.dependency_overrides[get_semantic_analyzer] = lambda: FakeAnalyzer()
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    _setup_profile(client, headers)
    application = _create_application(client, headers, "https://example.com/job/1")

    client.post(
        f"/applications/{application['id']}/interviews",
        headers=headers,
        json={"scheduled_at": "2026-09-10T09:00:00", "interview_type": "rh"},
    )
    client.post(
        f"/applications/{application['id']}/interviews",
        headers=headers,
        json={"scheduled_at": "2026-10-02T09:00:00", "interview_type": "manager"},
    )

    september = client.get(
        "/dashboard/calendar", headers=headers, params={"month": "2026-09"}
    )
    assert september.status_code == 200
    entries = september.json()
    assert len(entries) == 1
    assert entries[0]["interview_type"] == "rh"
    assert entries[0]["company_name"] == "Acme"

    invalid = client.get(
        "/dashboard/calendar", headers=headers, params={"month": "not-a-month"}
    )
    assert invalid.status_code == 422

    app.dependency_overrides.pop(get_semantic_analyzer, None)
