import io

from docx import Document

from app.llm.switch import set_llm_features_enabled
from app.llm_analyzer.analyzer import SemanticReport
from app.llm_analyzer.dependencies import get_semantic_analyzer
from app.main import app
from app.models.llm_call_log import LlmCallLog
from app.models.user import User
from tests._helpers import register_and_login


class _FakeAnalyzer:
    def analyze(self, cv_text: str, offer_text: str) -> SemanticReport:
        return SemanticReport(score=60, missing_keywords=[], recommendations=[])


def _cv_bytes() -> bytes:
    d = Document()
    d.add_paragraph("Expérience professionnelle")
    d.add_paragraph("Développeur")
    d.add_paragraph("Formation")
    d.add_paragraph("Master")
    d.add_paragraph("Compétences")
    d.add_paragraph("Python")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _post_diagnostic(client, token):
    return client.post(
        "/diagnostics",
        headers={"Authorization": f"Bearer {token}"},
        files={
            "cv_file": (
                "cv.docx",
                _cv_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"offer_text": "Python developer."},
    )


def _setup(client):
    app.dependency_overrides[get_semantic_analyzer] = lambda: _FakeAnalyzer()
    return register_and_login(client)


def _teardown():
    app.dependency_overrides.pop(get_semantic_analyzer, None)


def test_diagnostic_writes_one_llm_call_log(client, db_session):
    token = _setup(client)
    try:
        assert _post_diagnostic(client, token).status_code == 201
        rows = db_session.query(LlmCallLog).filter_by(feature="diagnostic").all()
        assert len(rows) == 1
    finally:
        _teardown()


def test_diagnostic_429_when_monthly_quota_reached(client, db_session):
    token = _setup(client)
    try:
        db_session.query(User).update({"quota_overrides": {"diagnostic": 1}})
        db_session.commit()
        assert _post_diagnostic(client, token).status_code == 201
        resp = _post_diagnostic(client, token)
        assert resp.status_code == 429
        assert resp.json()["detail"]["code"] == "quota_exceeded"
    finally:
        _teardown()


def test_diagnostic_503_when_llm_disabled(client, db_session):
    token = _setup(client)
    try:
        set_llm_features_enabled(db_session, False)
        resp = _post_diagnostic(client, token)
        assert resp.status_code == 503
        assert resp.json()["detail"]["code"] == "llm_paused"
    finally:
        set_llm_features_enabled(db_session, True)
        _teardown()
