import io

from docx import Document

from app.llm_analyzer.analyzer import SemanticReport
from app.llm_analyzer.dependencies import get_semantic_analyzer
from app.main import app

# ftp:// is rejected by app.offer_ingestion.scraper._validate_url on the
# scheme check alone, before any DNS/network I/O - keeps the background
# full_offer_text backfill (which TestClient runs synchronously) fast and
# network-free in tests.
_OFFER_URL = "ftp://example.com/job/1"


def _register_and_login(client, email: str = "jane@example.com") -> str:
    client.post("/auth/register", json={"email": email, "password": "s3cret!1"})
    login = client.post("/auth/login", data={"username": email, "password": "s3cret!1"})
    return login.json()["access_token"]


def _payload(**overrides) -> dict:
    payload = {
        "offer_url": _OFFER_URL,
        "title": "Développeur Python",
        "company": "Acme",
        "location": "Paris",
        "snippet": "Poste de développeur.",
        "source": "adzuna",
        "ats_type": None,
        "salary": None,
    }
    payload.update(overrides)
    return payload


def test_open_saved_job_creates_then_upserts_idempotently(client):
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}

    first = client.post("/saved-jobs", headers=headers, json=_payload())
    assert first.status_code == 200
    saved_id = first.json()["id"]

    second = client.post(
        "/saved-jobs", headers=headers, json=_payload(title="Développeur Python Senior")
    )
    assert second.status_code == 200
    assert second.json()["id"] == saved_id
    assert second.json()["title"] == "Développeur Python Senior"

    listed = client.get("/saved-jobs", headers=headers)
    assert len(listed.json()) == 1


def test_get_saved_job_returns_404_for_another_users_job(client):
    token_a = _register_and_login(client, "a@example.com")
    token_b = _register_and_login(client, "b@example.com")

    created = client.post(
        "/saved-jobs",
        headers={"Authorization": f"Bearer {token_a}"},
        json=_payload(),
    )
    saved_id = created.json()["id"]

    response = client.get(
        f"/saved-jobs/{saved_id}", headers={"Authorization": f"Bearer {token_b}"}
    )
    assert response.status_code == 404


def test_saved_jobs_require_auth(client):
    assert client.get("/saved-jobs").status_code == 401
    assert client.post("/saved-jobs", json=_payload()).status_code == 401


class FakeAnalyzer:
    def analyze(self, cv_text: str, offer_text: str) -> SemanticReport:
        return SemanticReport(score=70, missing_keywords=[], recommendations=[])


def _clean_cv_docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("Expérience professionnelle")
    document.add_paragraph("Développeur")
    document.add_paragraph("Formation")
    document.add_paragraph("Master")
    document.add_paragraph("Compétences")
    document.add_paragraph("Python")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_diagnostic_linked_to_saved_job_shows_up_in_detail(client):
    app.dependency_overrides[get_semantic_analyzer] = lambda: FakeAnalyzer()
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}

    saved_id = client.post("/saved-jobs", headers=headers, json=_payload()).json()["id"]

    diag_response = client.post(
        "/diagnostics",
        headers=headers,
        files={
            "cv_file": (
                "cv.docx",
                _clean_cv_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={
            "offer_text": "We need a Python developer.",
            "saved_job_id": str(saved_id),
        },
    )
    assert diag_response.status_code == 201

    detail = client.get(f"/saved-jobs/{saved_id}", headers=headers)
    assert detail.status_code == 200
    body = detail.json()
    assert body["latest_diagnostic"] is not None
    assert body["latest_diagnostic"]["overall_score"] > 0
    assert body["application_status"] is None

    app.dependency_overrides.pop(get_semantic_analyzer, None)


def test_create_diagnostic_rejects_saved_job_owned_by_another_user(client):
    app.dependency_overrides[get_semantic_analyzer] = lambda: FakeAnalyzer()
    token_a = _register_and_login(client, "a@example.com")
    token_b = _register_and_login(client, "b@example.com")

    saved_id = client.post(
        "/saved-jobs",
        headers={"Authorization": f"Bearer {token_a}"},
        json=_payload(),
    ).json()["id"]

    response = client.post(
        "/diagnostics",
        headers={"Authorization": f"Bearer {token_b}"},
        files={
            "cv_file": (
                "cv.docx",
                _clean_cv_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={
            "offer_text": "We need a Python developer.",
            "saved_job_id": str(saved_id),
        },
    )
    assert response.status_code == 404

    app.dependency_overrides.pop(get_semantic_analyzer, None)


_RENDER_PREVIEW_CONTENT = {
    "summary": "Développeuse Python expérimentée.",
    "experience": [
        {
            "title": "Développeuse",
            "company": "Acme",
            "dates": "2020-2022",
            "bullets": ["A conçu des API."],
        }
    ],
    "education": ["Master Informatique"],
    "skills": ["Python"],
}


def test_render_cv_preview_returns_pdf_without_any_llm_call(client):
    # No get_semantic_analyzer/get_cv_rewriter override is set up anywhere
    # in this test - if render-preview called an LLM under the hood, this
    # would hang or error trying to reach a real Anthropic client with the
    # test suite's fake API key. It doesn't: this endpoint is a pure
    # fpdf2 re-render from already-generated content.
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    saved_id = client.post("/saved-jobs", headers=headers, json=_payload()).json()["id"]

    response = client.post(
        f"/saved-jobs/{saved_id}/cv/render-preview",
        headers=headers,
        json={"content": _RENDER_PREVIEW_CONTENT, "template": "modern"},
    )
    assert response.status_code == 200
    assert response.headers["content-type"] == "application/pdf"
    assert response.content.startswith(b"%PDF")


def test_render_cv_preview_404s_for_another_users_saved_job(client):
    owner_token = _register_and_login(client, "a@example.com")
    saved_id = client.post(
        "/saved-jobs",
        headers={"Authorization": f"Bearer {owner_token}"},
        json=_payload(),
    ).json()["id"]

    attacker_token = _register_and_login(client, "b@example.com")
    response = client.post(
        f"/saved-jobs/{saved_id}/cv/render-preview",
        headers={"Authorization": f"Bearer {attacker_token}"},
        json={"content": _RENDER_PREVIEW_CONTENT},
    )
    assert response.status_code == 404
