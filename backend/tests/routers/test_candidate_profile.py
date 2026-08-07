import io

from docx import Document

from app.main import app


def _register_and_login(client, email: str = "jane@example.com") -> str:
    client.post("/auth/register", json={"email": email, "password": "s3cret!1"})
    login = client.post("/auth/login", data={"username": email, "password": "s3cret!1"})
    return login.json()["access_token"]


def _clean_cv_docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("Expérience professionnelle")
    document.add_paragraph("Développeuse Full Stack chez Acme, 2020-2022")
    document.add_paragraph("Formation")
    document.add_paragraph("Master Informatique")
    document.add_paragraph("Compétences")
    document.add_paragraph("Python, Docker")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_get_profile_returns_404_when_not_yet_created(client):
    token = _register_and_login(client)
    response = client.get("/profile", headers={"Authorization": f"Bearer {token}"})
    assert response.status_code == 404


def test_put_profile_creates_then_updates(client):
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}

    create = client.put(
        "/profile",
        headers=headers,
        json={
            "full_name": "Jane Doe",
            "phone": "0612345678",
            "work_authorization": "Autorisée à travailler en France/UE",
        },
    )
    assert create.status_code == 200
    assert create.json()["full_name"] == "Jane Doe"
    assert create.json()["has_cv"] is False

    update = client.put(
        "/profile",
        headers=headers,
        json={
            "full_name": "Jane A. Doe",
            "phone": "0612345678",
            "work_authorization": "Autorisée à travailler en France/UE",
            "salary_expectation": "45-55k€",
        },
    )
    assert update.status_code == 200
    assert update.json()["full_name"] == "Jane A. Doe"
    assert update.json()["salary_expectation"] == "45-55k€"

    fetched = client.get("/profile", headers=headers)
    assert fetched.json()["full_name"] == "Jane A. Doe"


def test_upload_cv_parses_and_stores_reference_cv(client):
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    client.put(
        "/profile",
        headers=headers,
        json={"full_name": "Jane Doe", "phone": "0612345678", "work_authorization": "FR/UE"},
    )

    response = client.post(
        "/profile/cv",
        headers=headers,
        files={
            "cv_file": (
                "cv.docx",
                _clean_cv_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    body = response.json()
    assert body["has_cv"] is True
    assert body["cv_filename"] == "cv.docx"


def test_upload_cv_rejects_unsupported_format(client):
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    client.put(
        "/profile",
        headers=headers,
        json={"full_name": "Jane Doe", "phone": "0612345678", "work_authorization": "FR/UE"},
    )

    response = client.post(
        "/profile/cv",
        headers=headers,
        files={"cv_file": ("cv.txt", b"plain text resume", "text/plain")},
    )

    assert response.status_code == 422


def test_upload_cv_before_put_creates_profile_implicitly(client):
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}

    response = client.post(
        "/profile/cv",
        headers=headers,
        files={
            "cv_file": (
                "cv.docx",
                _clean_cv_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    assert response.json()["has_cv"] is True
    assert response.json()["full_name"] == ""


def test_profile_endpoints_require_auth(client):
    assert client.get("/profile").status_code == 401
    assert client.put("/profile", json={"full_name": "x", "phone": "x", "work_authorization": "x"}).status_code == 401


def test_delete_profile_removes_the_stored_profile(client):
    # Finding 4 (RGPD): the CandidateProfile holds the most personal data in
    # the system (full reference CV text, phone, address) and had no
    # deletion path at all - DELETE /diagnostics never touched it.
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    client.put(
        "/profile",
        headers=headers,
        json={"full_name": "Jane Doe", "phone": "0612345678", "work_authorization": "FR/UE"},
    )
    client.post(
        "/profile/cv",
        headers=headers,
        files={
            "cv_file": (
                "cv.docx",
                _clean_cv_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    response = client.delete("/profile", headers=headers)

    assert response.status_code == 204
    assert client.get("/profile", headers=headers).status_code == 404


def test_delete_profile_is_idempotent_when_no_profile_exists(client):
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}

    assert client.delete("/profile", headers=headers).status_code == 204
    assert client.delete("/profile", headers=headers).status_code == 204


def test_delete_profile_only_deletes_the_current_users_profile(client):
    other_headers = {"Authorization": f"Bearer {_register_and_login(client, email='other@example.com')}"}
    client.put(
        "/profile",
        headers=other_headers,
        json={"full_name": "Other User", "phone": "0600000000", "work_authorization": "FR/UE"},
    )
    headers = {"Authorization": f"Bearer {_register_and_login(client)}"}
    client.put(
        "/profile",
        headers=headers,
        json={"full_name": "Jane Doe", "phone": "0612345678", "work_authorization": "FR/UE"},
    )

    assert client.delete("/profile", headers=headers).status_code == 204

    assert client.get("/profile", headers=other_headers).status_code == 200


def test_delete_profile_requires_auth(client):
    assert client.delete("/profile").status_code == 401
