import io

from docx import Document

from app.llm_analyzer.analyzer import SemanticReport
from app.llm_analyzer.dependencies import get_semantic_analyzer
from app.main import app
from app.personalization.dependencies import get_cover_letter_generator, get_cv_rewriter
from app.personalization.schemas import CoverLetter, CvExperienceEntry, RewrittenCv
from app.rate_limit.limiter import MAX_PERSONALIZATIONS_PER_HOUR
from app.storage.client import ObjectStorage, ObjectStorageError
from app.storage.dependencies import get_object_storage


class FakeSemanticAnalyzer:
    """Used for the ats_score_after recompute inside run_cv_generation_job
    - distinct from the FakeAnalyzer used to create the diagnostic itself,
    kept around for the whole /cv call unlike that one (see
    _create_diagnostic, which pops its own override immediately after
    creating the diagnostic)."""

    def analyze(self, cv_text: str, offer_text: str) -> SemanticReport:
        return SemanticReport(score=75, missing_keywords=[], recommendations=[])


class FakeCvRewriter:
    def rewrite(
        self,
        cv_text,
        offer_text,
        missing_keywords,
        recommendations,
        stricter_length=False,
        **kwargs,
    ):
        return RewrittenCv(
            summary="Résumé optimisé.",
            experience=[
                CvExperienceEntry(
                    title="Développeuse",
                    company="Acme",
                    dates="2020-2022",
                    bullets=["A conçu des API."],
                )
            ],
            education=["Master Informatique"],
            skills=["Python"],
        )


class FailingCvRewriter:
    def rewrite(
        self,
        cv_text,
        offer_text,
        missing_keywords,
        recommendations,
        stricter_length=False,
        **kwargs,
    ):
        from app.personalization.analyzer import PersonalizationError

        raise PersonalizationError("boom")


class OverflowingThenShortCvRewriter:
    """First call returns a CV long enough to overflow a single A4 page;
    the retry call (stricter_length=True) returns a short one - exercises
    the router's post-generation page-count retry."""

    def __init__(self):
        self.stricter_length_flags: list[bool] = []

    def rewrite(
        self,
        cv_text,
        offer_text,
        missing_keywords,
        recommendations,
        stricter_length=False,
        **kwargs,
    ):
        self.stricter_length_flags.append(stricter_length)
        if not stricter_length:
            return RewrittenCv(
                summary="Résumé optimisé. " * 30,
                experience=[
                    CvExperienceEntry(
                        title="Développeuse",
                        company="Acme",
                        dates="2020-2022",
                        bullets=["A conçu des API performantes et robustes. " * 8] * 6,
                    )
                    for _ in range(8)
                ],
                education=["Master Informatique"] * 5,
                skills=["Python"] * 30,
            )
        return RewrittenCv(
            summary="Résumé optimisé, version courte.",
            experience=[
                CvExperienceEntry(
                    title="Développeuse",
                    company="Acme",
                    dates="2020-2022",
                    bullets=["A conçu des API."],
                )
            ],
            education=["Master Informatique"],
            skills=["Python"],
        )


class FakeCoverLetterGenerator:
    def __init__(self):
        self.tones_seen: list[str] = []

    def generate(
        self, cv_text, offer_text, missing_keywords, recommendations, tone="sobre"
    ):
        self.tones_seen.append(tone)
        return CoverLetter(
            greeting="Madame, Monsieur,",
            body_paragraphs=[f"Je vous écris pour candidater à ce poste ({tone})."],
            closing_formula="Cordialement,",
            signature="Jane Doe",
        )


class FailingCoverLetterGenerator:
    def generate(
        self, cv_text, offer_text, missing_keywords, recommendations, tone="sobre"
    ):
        from app.personalization.analyzer import PersonalizationError

        raise PersonalizationError("boom")


class FakeObjectStorage(ObjectStorage):
    def __init__(self):
        self._objects: dict[str, bytes] = {}

    def upload(self, key: str, content: bytes) -> None:
        self._objects[key] = content

    def download(self, key: str) -> bytes:
        if key not in self._objects:
            raise ObjectStorageError(f"missing key {key}")
        return self._objects[key]

    def delete(self, key: str) -> None:
        self._objects.pop(key, None)


def _clean_cv_docx_bytes() -> bytes:
    # Includes "Acme", "2020-2022", and "Master Informatique" verbatim so
    # that FakeCvRewriter's output below (which reuses these exact terms)
    # doesn't trip app.personalization.verification.cv_needs_review's
    # anti-hallucination check: that check flags any year or multi-word
    # capitalized phrase in the rewritten CV that isn't already present in
    # the original CV text, and correctly treats made-up employers/dates/
    # diplomas as needing review (see tests/personalization/test_verification.py).
    document = Document()
    document.add_paragraph("Expérience professionnelle")
    document.add_paragraph("Développeur chez Acme, 2020-2022")
    document.add_paragraph("Formation")
    document.add_paragraph("Master Informatique")
    document.add_paragraph("Compétences")
    document.add_paragraph("Python")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _register_and_login(client, email: str = "jane@example.com") -> str:
    client.post("/auth/register", json={"email": email, "password": "s3cret!1"})
    login = client.post("/auth/login", data={"username": email, "password": "s3cret!1"})
    return login.json()["access_token"]


def _create_diagnostic(client, headers) -> int:
    from app.llm_analyzer.analyzer import SemanticReport
    from app.llm_analyzer.dependencies import get_semantic_analyzer

    class FakeAnalyzer:
        def analyze(self, cv_text: str, offer_text: str) -> SemanticReport:
            return SemanticReport(
                score=60, missing_keywords=["Docker"], recommendations=["Add Docker"]
            )

    app.dependency_overrides[get_semantic_analyzer] = lambda: FakeAnalyzer()
    response = client.post(
        "/diagnostics",
        headers=headers,
        files={
            "cv_file": (
                "cv.docx",
                _clean_cv_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"offer_text": "We need a Python developer with Docker experience."},
    )
    app.dependency_overrides.pop(get_semantic_analyzer, None)
    return response.json()["id"]


def _override_personalization_deps():
    # A single shared FakeObjectStorage instance, not `lambda:
    # FakeObjectStorage()`. FastAPI calls a dependency_overrides callable
    # fresh on every request, so `lambda: FakeObjectStorage()` would hand
    # each request (the POST that uploads, then the GET that downloads) its
    # own empty instance and every download would 404/503 with "missing
    # key" - the real get_object_storage dependency is @lru_cache'd to a
    # singleton in production, so the fake must behave the same way for the
    # life of a test.
    storage = FakeObjectStorage()
    app.dependency_overrides[get_cv_rewriter] = lambda: FakeCvRewriter()
    app.dependency_overrides[get_cover_letter_generator] = lambda: (
        FakeCoverLetterGenerator()
    )
    app.dependency_overrides[get_object_storage] = lambda: storage
    app.dependency_overrides[get_semantic_analyzer] = lambda: FakeSemanticAnalyzer()


def _clear_personalization_overrides():
    app.dependency_overrides.pop(get_cv_rewriter, None)
    app.dependency_overrides.pop(get_cover_letter_generator, None)
    app.dependency_overrides.pop(get_object_storage, None)
    app.dependency_overrides.pop(get_semantic_analyzer, None)


def _generate_cv_and_wait(client, headers, diagnostic_id, **form_data):
    """POST /diagnostics/{id}/cv now launches a background job (202 + job_id)
    instead of generating synchronously. TestClient runs BackgroundTasks
    synchronously, so by the time this returns the job has already finished
    - this just polls once to fetch the terminal GenerationJobOut."""
    launch = client.post(
        f"/diagnostics/{diagnostic_id}/cv", headers=headers, data=form_data
    )
    assert launch.status_code == 202
    job_id = launch.json()["job_id"]
    job = client.get(f"/generation-jobs/{job_id}", headers=headers)
    assert job.status_code == 200
    return job.json()


def _generate_lettre_and_wait(client, headers, diagnostic_id, **form_data):
    """POST /diagnostics/{id}/lettre launches a background job (202 +
    job_id), same as the CV endpoint - see _generate_cv_and_wait."""
    launch = client.post(
        f"/diagnostics/{diagnostic_id}/lettre", headers=headers, data=form_data
    )
    assert launch.status_code == 202
    job_id = launch.json()["job_id"]
    job = client.get(f"/generation-jobs/{job_id}", headers=headers)
    assert job.status_code == 200
    return job.json()


def test_generate_cv_returns_metadata_and_download_serves_pdf(client):
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    diagnostic_id = _create_diagnostic(client, headers)
    _override_personalization_deps()

    job = _generate_cv_and_wait(client, headers, diagnostic_id)
    assert job["status"] == "done"
    assert job["result"]["kind"] == "cv"
    assert job["result"]["needs_review"] is False
    assert job["result"]["ats_score_before"] is not None
    assert job["result"]["ats_score_after"] is not None

    download = client.get(f"/diagnostics/{diagnostic_id}/cv", headers=headers)
    assert download.status_code == 200
    assert download.headers["content-type"] == "application/pdf"
    assert download.content.startswith(b"%PDF")

    _clear_personalization_overrides()


def test_generate_cv_retries_with_stricter_prompt_when_pdf_overflows_one_page(client):
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    diagnostic_id = _create_diagnostic(client, headers)
    _override_personalization_deps()
    rewriter = OverflowingThenShortCvRewriter()
    app.dependency_overrides[get_cv_rewriter] = lambda: rewriter

    job = _generate_cv_and_wait(client, headers, diagnostic_id)
    assert job["status"] == "done"

    # First call with the normal prompt, second (retry) with stricter_length=True.
    assert rewriter.stricter_length_flags == [False, True]

    download = client.get(f"/diagnostics/{diagnostic_id}/cv", headers=headers)
    assert download.status_code == 200
    import pdfplumber

    with pdfplumber.open(io.BytesIO(download.content)) as pdf:
        assert len(pdf.pages) == 1

    _clear_personalization_overrides()


def test_generate_lettre_returns_metadata_and_download_serves_pdf(client):
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    diagnostic_id = _create_diagnostic(client, headers)
    _override_personalization_deps()

    job = _generate_lettre_and_wait(client, headers, diagnostic_id)
    assert job["status"] == "done"
    assert job["result"]["kind"] == "lettre"
    assert job["result"]["needs_review"] is False

    download = client.get(f"/diagnostics/{diagnostic_id}/lettre", headers=headers)
    assert download.status_code == 200
    assert download.content.startswith(b"%PDF")

    _clear_personalization_overrides()


def test_generate_lettre_passes_tone_through_to_generator(client):
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    diagnostic_id = _create_diagnostic(client, headers)
    _override_personalization_deps()
    generator = FakeCoverLetterGenerator()
    app.dependency_overrides[get_cover_letter_generator] = lambda: generator

    job = _generate_lettre_and_wait(client, headers, diagnostic_id, tone="direct")
    assert job["status"] == "done"
    assert generator.tones_seen == ["direct"]

    _clear_personalization_overrides()


def test_generate_lettre_job_ends_in_error_status_on_llm_failure(client):
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    diagnostic_id = _create_diagnostic(client, headers)
    _override_personalization_deps()
    app.dependency_overrides[get_cover_letter_generator] = lambda: (
        FailingCoverLetterGenerator()
    )

    job = _generate_lettre_and_wait(client, headers, diagnostic_id)
    assert job["status"] == "error"
    assert job["error"]

    _clear_personalization_overrides()


def test_regenerating_cv_replaces_the_previous_document(client):
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    diagnostic_id = _create_diagnostic(client, headers)
    _override_personalization_deps()

    first = _generate_cv_and_wait(client, headers, diagnostic_id)["result"]
    second = _generate_cv_and_wait(client, headers, diagnostic_id)["result"]

    assert first["created_at"] == second["created_at"]
    assert second["updated_at"] >= first["updated_at"]

    _clear_personalization_overrides()


def test_generate_cv_for_missing_diagnostic_returns_404(client):
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    _override_personalization_deps()

    response = client.post("/diagnostics/999999/cv", headers=headers)
    assert response.status_code == 404

    _clear_personalization_overrides()


def test_cannot_generate_or_download_cv_for_another_users_diagnostic(client):
    owner_token = _register_and_login(client, "jane@example.com")
    owner_headers = {"Authorization": f"Bearer {owner_token}"}
    diagnostic_id = _create_diagnostic(client, owner_headers)
    _override_personalization_deps()
    client.post(f"/diagnostics/{diagnostic_id}/cv", headers=owner_headers)

    attacker_token = _register_and_login(client, "mallory@example.com")
    attacker_headers = {"Authorization": f"Bearer {attacker_token}"}

    generate = client.post(f"/diagnostics/{diagnostic_id}/cv", headers=attacker_headers)
    assert generate.status_code == 404

    download = client.get(f"/diagnostics/{diagnostic_id}/cv", headers=attacker_headers)
    assert download.status_code == 404

    _clear_personalization_overrides()


def test_download_cv_before_generation_returns_404(client):
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    diagnostic_id = _create_diagnostic(client, headers)
    _override_personalization_deps()

    response = client.get(f"/diagnostics/{diagnostic_id}/cv", headers=headers)
    assert response.status_code == 404

    _clear_personalization_overrides()


def test_generate_cv_job_ends_in_error_status_on_llm_failure(client):
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    diagnostic_id = _create_diagnostic(client, headers)
    _override_personalization_deps()
    app.dependency_overrides[get_cv_rewriter] = lambda: FailingCvRewriter()

    job = _generate_cv_and_wait(client, headers, diagnostic_id)
    assert job["status"] == "error"
    assert job["error"]

    _clear_personalization_overrides()


def test_personalization_rate_limit_returns_429(client):
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    diagnostic_id = _create_diagnostic(client, headers)
    _override_personalization_deps()

    for _ in range(MAX_PERSONALIZATIONS_PER_HOUR):
        response = client.post(f"/diagnostics/{diagnostic_id}/cv", headers=headers)
        assert response.status_code == 202

    blocked = client.post(f"/diagnostics/{diagnostic_id}/cv", headers=headers)
    assert blocked.status_code == 429

    _clear_personalization_overrides()


def test_personalization_request_log_only_written_on_job_success(client, db_session):
    from app.models.personalization_request_log import PersonalizationRequestLog

    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    diagnostic_id = _create_diagnostic(client, headers)
    _override_personalization_deps()
    app.dependency_overrides[get_cv_rewriter] = lambda: FailingCvRewriter()

    def _log_count() -> int:
        return db_session.query(PersonalizationRequestLog).count()

    before = _log_count()
    failed_job = _generate_cv_and_wait(client, headers, diagnostic_id)
    assert failed_job["status"] == "error"
    assert _log_count() == before

    app.dependency_overrides[get_cv_rewriter] = lambda: FakeCvRewriter()
    succeeded_job = _generate_cv_and_wait(client, headers, diagnostic_id)
    assert succeeded_job["status"] == "done"
    assert _log_count() == before + 1

    _clear_personalization_overrides()
