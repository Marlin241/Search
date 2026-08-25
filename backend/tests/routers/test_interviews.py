from app.llm_analyzer.analyzer import SemanticReport
from app.llm_analyzer.dependencies import get_semantic_analyzer
from app.main import app


class FakeAnalyzer:
    def analyze(self, cv_text, offer_text):
        return SemanticReport(
            score=70, missing_keywords=["Docker"], recommendations=["Add Docker"]
        )


def _register_and_login(client, email: str = "jane@example.com") -> str:
    client.post("/auth/register", json={"email": email, "password": "s3cret!1"})
    login = client.post("/auth/login", data={"username": email, "password": "s3cret!1"})
    return login.json()["access_token"]


def _setup_profile(client, headers: dict) -> None:
    client.put(
        "/profile",
        headers=headers,
        json={
            "full_name": "Jane Doe",
            "phone": "0600000000",
            "work_authorization": "FR/UE",
        },
    )
    import io

    from docx import Document

    document = Document()
    document.add_paragraph("Expérience professionnelle")
    document.add_paragraph("Développeuse Full Stack chez Acme, 2020-2022")
    document.add_paragraph("Formation")
    document.add_paragraph("Master Informatique")
    document.add_paragraph("Compétences")
    document.add_paragraph("Python, Docker")
    buffer = io.BytesIO()
    document.save(buffer)
    client.post(
        "/profile/cv",
        headers=headers,
        files={
            "cv_file": (
                "cv.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )


def _create_application(client, headers: dict, offer_url: str) -> dict:
    response = client.post(
        "/applications",
        headers=headers,
        json={
            "offer_url": offer_url,
            "offer_text": "Offre.",
            "source": "manual",
            "company_name": "Acme",
            "job_title": "Dev",
        },
    )
    assert response.status_code == 201
    return response.json()


def test_create_list_update_delete_interview(client):
    app.dependency_overrides[get_semantic_analyzer] = lambda: FakeAnalyzer()
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    _setup_profile(client, headers)
    application = _create_application(client, headers, "https://example.com/job/1")

    created = client.post(
        f"/applications/{application['id']}/interviews",
        headers=headers,
        json={
            "scheduled_at": "2026-09-15T10:00:00",
            "interview_type": "rh",
            "location_or_link": "https://meet.example.com/x",
            "notes": "Premier échange",
        },
    )
    assert created.status_code == 201
    body = created.json()
    assert body["application_id"] == application["id"]
    assert body["interview_type"] == "rh"
    interview_id = body["id"]

    listed = client.get(
        f"/applications/{application['id']}/interviews", headers=headers
    )
    assert listed.status_code == 200
    assert len(listed.json()) == 1
    assert listed.json()[0]["id"] == interview_id

    updated = client.patch(
        f"/interviews/{interview_id}",
        headers=headers,
        json={"interview_type": "manager", "notes": "Reprogrammé"},
    )
    assert updated.status_code == 200
    assert updated.json()["interview_type"] == "manager"
    assert updated.json()["notes"] == "Reprogrammé"
    # Fields not included in the PATCH body are left untouched.
    assert updated.json()["location_or_link"] == "https://meet.example.com/x"

    deleted = client.delete(f"/interviews/{interview_id}", headers=headers)
    assert deleted.status_code == 204
    assert (
        client.get(
            f"/applications/{application['id']}/interviews", headers=headers
        ).json()
        == []
    )

    app.dependency_overrides.pop(get_semantic_analyzer, None)


def test_interview_endpoints_reject_another_users_application(client):
    app.dependency_overrides[get_semantic_analyzer] = lambda: FakeAnalyzer()
    owner_token = _register_and_login(client, "owner@example.com")
    owner_headers = {"Authorization": f"Bearer {owner_token}"}
    _setup_profile(client, owner_headers)
    application = _create_application(
        client, owner_headers, "https://example.com/job/1"
    )

    attacker_token = _register_and_login(client, "attacker@example.com")
    attacker_headers = {"Authorization": f"Bearer {attacker_token}"}

    assert (
        client.post(
            f"/applications/{application['id']}/interviews",
            headers=attacker_headers,
            json={"scheduled_at": "2026-09-15T10:00:00", "interview_type": "rh"},
        ).status_code
        == 404
    )
    assert (
        client.get(
            f"/applications/{application['id']}/interviews", headers=attacker_headers
        ).status_code
        == 404
    )

    app.dependency_overrides.pop(get_semantic_analyzer, None)


def test_interview_not_found_returns_404(client):
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    assert (
        client.patch(
            "/interviews/999999",
            headers=headers,
            json={"interview_type": "jury"},
        ).status_code
        == 404
    )
    assert client.delete("/interviews/999999", headers=headers).status_code == 404
