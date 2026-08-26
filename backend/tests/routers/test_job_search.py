import io

import httpx
import respx
from docx import Document

from app.compatibility.analyzer import CompatibilityAnalysisError, CompatibilityDetail
from app.compatibility.dependencies import get_compatibility_detail_analyzer
from app.job_search.dependencies import get_job_search_clients
from app.job_search.errors import JobSearchSourceError
from app.job_search.greenhouse import GreenhouseJobBoardClient
from app.job_search.lever import LeverJobBoardClient
from app.job_search.schemas import JobListing
from app.main import app
from app.rate_limit.limiter import (
    MAX_COMPATIBILITY_DETAILS_PER_HOUR,
    MAX_SEARCHES_PER_HOUR,
)


def _register_and_login(client, email: str = "jane@example.com") -> str:
    client.post("/auth/register", json={"email": email, "password": "s3cret!1"})
    login = client.post("/auth/login", data={"username": email, "password": "s3cret!1"})
    return login.json()["access_token"]


class FakeWorkingClient:
    """Used by tests that don't care about company discovery. `company` is
    deliberately blank — extract_unique_companies() skips blank names — so
    these tests never trigger the background discovery path (which would
    otherwise make real, unmocked HTTP calls to Greenhouse/Lever). Tests that
    DO want to exercise discovery use CompanyMentioningClient below."""

    def search(self, criteria):
        return [
            JobListing(
                title="Développeur Python",
                company="",
                location="Paris",
                snippet="...",
                url="https://example.com/1",
                source="fake",
                ats_type=None,
            )
        ]


class CompanyMentioningClient:
    def search(self, criteria):
        return [
            JobListing(
                title="Développeur Python",
                company="Acme",
                location="Paris",
                snippet="...",
                url="https://example.com/1",
                source="fake",
                ats_type=None,
            )
        ]


class FakeFailingClient:
    def search(self, criteria):
        raise JobSearchSourceError("down")


class EmptyGreenhouseOrLeverClient:
    def search(self, criteria, company_slugs):
        return []


class EmptyPrimaryClient:
    def search(self, criteria):
        return []


def _default_clients(overrides: dict[str, object]) -> dict[str, object]:
    base: dict[str, object] = {
        "france_travail": EmptyPrimaryClient(),
        "adzuna": EmptyPrimaryClient(),
        "la_bonne_alternance": EmptyPrimaryClient(),
        "greenhouse": EmptyGreenhouseOrLeverClient(),
        "lever": EmptyGreenhouseOrLeverClient(),
    }
    base.update(overrides)
    return base


def test_search_returns_listings_and_unavailable_sources(client):
    app.dependency_overrides[get_job_search_clients] = lambda: _default_clients(
        {"france_travail": FakeWorkingClient(), "adzuna": FakeFailingClient()}
    )
    token = _register_and_login(client)

    response = client.post(
        "/job-search/search",
        headers={"Authorization": f"Bearer {token}"},
        json={"keywords": "python"},
    )

    assert response.status_code == 200
    body = response.json()
    assert len(body["listings"]) == 1
    assert body["unavailable_sources"] == ["adzuna"]
    assert "search_id" in body


def test_search_requires_auth(client):
    app.dependency_overrides[get_job_search_clients] = lambda: _default_clients(
        {"france_travail": FakeWorkingClient()}
    )
    response = client.post("/job-search/search", json={"keywords": "python"})
    assert response.status_code == 401


def test_search_rate_limited_after_max_per_hour(client):
    app.dependency_overrides[get_job_search_clients] = lambda: _default_clients(
        {"france_travail": FakeWorkingClient()}
    )
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}

    for _ in range(MAX_SEARCHES_PER_HOUR):
        response = client.post(
            "/job-search/search", headers=headers, json={"keywords": "python"}
        )
        assert response.status_code == 200

    response = client.post(
        "/job-search/search", headers=headers, json={"keywords": "python"}
    )
    assert response.status_code == 429


def test_search_with_no_companies_in_results_is_not_discovery_pending(client):
    class NoCompanyClient:
        def search(self, criteria):
            return []

    app.dependency_overrides[get_job_search_clients] = lambda: _default_clients(
        {"france_travail": NoCompanyClient()}
    )
    token = _register_and_login(client)

    response = client.post(
        "/job-search/search",
        headers={"Authorization": f"Bearer {token}"},
        json={"keywords": "python"},
    )

    assert response.json()["discovery_pending"] is False


@respx.mock
def test_search_discovers_unknown_company_and_polling_returns_new_listing(client):
    respx.get("https://boards-api.greenhouse.io/v1/boards/acme/jobs").mock(
        return_value=httpx.Response(
            200,
            json={
                "jobs": [
                    {
                        "title": "Ingénieur backend Python",
                        "location": {"name": "Paris"},
                        "content": "<p>Poste Acme.</p>",
                        "absolute_url": "https://boards.greenhouse.io/acme/jobs/1",
                    }
                ]
            },
        )
    )

    app.dependency_overrides[get_job_search_clients] = lambda: {
        "france_travail": CompanyMentioningClient(),
        "adzuna": EmptyPrimaryClient(),
        "la_bonne_alternance": EmptyPrimaryClient(),
        "greenhouse": GreenhouseJobBoardClient(),
        "lever": LeverJobBoardClient(),
    }
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}

    response = client.post(
        "/job-search/search", headers=headers, json={"keywords": "python"}
    )
    assert response.status_code == 200
    body = response.json()
    assert body["discovery_pending"] is True
    search_id = body["search_id"]

    poll = client.get(f"/job-search/search/{search_id}/discovery", headers=headers)
    assert poll.status_code == 200
    poll_body = poll.json()
    assert poll_body["done"] is True
    assert len(poll_body["new_listings"]) == 1
    assert poll_body["new_listings"][0]["title"] == "Ingénieur backend Python"


@respx.mock
def test_search_with_dakar_location_returns_waves_real_listing_synchronously(client):
    respx.get("https://boards-api.greenhouse.io/v1/boards/wavemm1/jobs").mock(
        return_value=httpx.Response(
            200,
            json={
                "jobs": [
                    {
                        "title": "Customer Support Team Lead",
                        "location": {"name": "Dakar, Senegal"},
                        "content": "<p>Poste Wave.</p>",
                        "absolute_url": "https://www.wave.com/en/careers/job/1",
                    }
                ]
            },
        )
    )
    respx.get(
        url__regex=r"https://boards-api\.greenhouse\.io/v1/boards/(?!wavemm1)[a-z0-9-]+/jobs"
    ).mock(return_value=httpx.Response(404))
    respx.get(url__regex=r"https://api\.lever\.co/v0/postings/[a-z0-9-]+").mock(
        return_value=httpx.Response(404)
    )

    app.dependency_overrides[get_job_search_clients] = lambda: {
        "france_travail": EmptyPrimaryClient(),
        "adzuna": EmptyPrimaryClient(),
        "la_bonne_alternance": EmptyPrimaryClient(),
        "greenhouse": GreenhouseJobBoardClient(),
        "lever": LeverJobBoardClient(),
    }
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}

    response = client.post(
        "/job-search/search",
        headers=headers,
        json={"keywords": "Support", "location": "Dakar"},
    )

    assert response.status_code == 200
    titles = [listing["title"] for listing in response.json()["listings"]]
    assert "Customer Support Team Lead" in titles


@respx.mock
def test_search_with_dakar_location_triggers_discovery_from_seed_companies_even_with_no_primary_results(
    client,
):
    respx.get(url__regex=r"https://boards-api\.greenhouse\.io/v1/boards/.+/jobs").mock(
        return_value=httpx.Response(404)
    )
    respx.get(url__regex=r"https://api\.lever\.co/v0/postings/.+").mock(
        return_value=httpx.Response(404)
    )

    app.dependency_overrides[get_job_search_clients] = lambda: _default_clients({})
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}

    response = client.post(
        "/job-search/search",
        headers=headers,
        json={"keywords": "python", "location": "Dakar"},
    )

    assert response.status_code == 200
    assert response.json()["discovery_pending"] is True


def test_get_discovery_for_unknown_search_id_returns_done_true(client):
    app.dependency_overrides[get_job_search_clients] = lambda: _default_clients(
        {"france_travail": FakeWorkingClient()}
    )
    token = _register_and_login(client)

    response = client.get(
        "/job-search/search/does-not-exist/discovery",
        headers={"Authorization": f"Bearer {token}"},
    )

    assert response.status_code == 200
    assert response.json() == {"done": True, "new_listings": []}


def test_get_discovery_requires_auth(client):
    response = client.get("/job-search/search/some-id/discovery")
    assert response.status_code == 401


class TwoListingsClient:
    """One listing matches the candidate's desired title, the other doesn't -
    used to assert the search endpoint scores and sorts by compatibility."""

    def search(self, criteria):
        return [
            JobListing(
                title="Comptable senior",
                company="",
                location="Paris",
                snippet="...",
                url="https://example.com/unrelated",
                source="fake",
                ats_type=None,
            ),
            JobListing(
                title="Développeur Python",
                company="",
                location="Paris",
                snippet="...",
                url="https://example.com/match",
                source="fake",
                ats_type=None,
            ),
        ]


def _clean_cv_docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("Expérience professionnelle")
    document.add_paragraph("Développeuse Full Stack chez Acme, 2020-2022")
    document.add_paragraph("Formation")
    document.add_paragraph("Master Informatique")
    document.add_paragraph("Compétences")
    document.add_paragraph("Python, Docker")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _onboard_with_cv(client, headers: dict[str, str]) -> None:
    client.put(
        "/profile",
        headers=headers,
        json={
            "first_name": "Jane",
            "last_name": "Doe",
            "phone": "0612345678",
            "work_authorization": "FR/UE",
        },
    )
    client.post(
        "/profile/cv",
        headers=headers,
        files={
            "cv_file": (
                "cv.docx",
                _clean_cv_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    client.put(
        "/profile/onboarding",
        headers=headers,
        json={
            "first_name": "Jane",
            "last_name": "Doe",
            "desired_job_titles": ["Développeur Python"],
        },
    )


def test_search_response_includes_compatibility_score(client):
    app.dependency_overrides[get_job_search_clients] = lambda: _default_clients(
        {"france_travail": FakeWorkingClient()}
    )
    token = _register_and_login(client)

    response = client.post(
        "/job-search/search",
        headers={"Authorization": f"Bearer {token}"},
        json={"keywords": "python"},
    )

    assert response.status_code == 200
    assert "compatibility_score" in response.json()["listings"][0]


def test_search_sorts_listings_by_compatibility_score_descending(client):
    app.dependency_overrides[get_job_search_clients] = lambda: _default_clients(
        {"france_travail": TwoListingsClient()}
    )
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    _onboard_with_cv(client, headers)

    response = client.post(
        "/job-search/search", headers=headers, json={"keywords": "python"}
    )

    listings = response.json()["listings"]
    assert listings[0]["title"] == "Développeur Python"
    assert listings[0]["compatibility_score"] >= listings[1]["compatibility_score"]


class FakeCompatibilityDetailAnalyzer:
    def analyze(self, cv_text, offer_text, score_breakdown):
        return CompatibilityDetail(
            summary="Bon profil pour ce poste.",
            strengths=["Expérience Python pertinente"],
            concerns=["Pas de mention de Docker dans l'offre"],
        )


class FailingCompatibilityDetailAnalyzer:
    def analyze(self, cv_text, offer_text, score_breakdown):
        raise CompatibilityAnalysisError("LLM indisponible")


def _sample_listing_payload() -> dict:
    return {
        "listing": {
            "title": "Développeur Python",
            "company": "Acme",
            "location": "Paris",
            "snippet": "Poste de développeur Python, 3 ans d'expérience.",
            "url": "https://example.com/job/1",
            "source": "adzuna",
            "ats_type": None,
        }
    }


def test_compatibility_detail_returns_breakdown_and_explanation(client):
    app.dependency_overrides[get_compatibility_detail_analyzer] = lambda: (
        FakeCompatibilityDetailAnalyzer()
    )
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    _onboard_with_cv(client, headers)

    response = client.post(
        "/job-search/compatibility-detail",
        headers=headers,
        json=_sample_listing_payload(),
    )

    assert response.status_code == 200
    body = response.json()
    assert body["summary"] == "Bon profil pour ce poste."
    assert "overall" in body["breakdown"]

    app.dependency_overrides.pop(get_compatibility_detail_analyzer, None)


def test_compatibility_detail_requires_cv(client):
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}

    response = client.post(
        "/job-search/compatibility-detail",
        headers=headers,
        json=_sample_listing_payload(),
    )

    assert response.status_code == 422


def test_compatibility_detail_requires_auth(client):
    response = client.post(
        "/job-search/compatibility-detail", json=_sample_listing_payload()
    )
    assert response.status_code == 401


def test_compatibility_detail_propagates_llm_failure_as_503(client):
    app.dependency_overrides[get_compatibility_detail_analyzer] = lambda: (
        FailingCompatibilityDetailAnalyzer()
    )
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    _onboard_with_cv(client, headers)

    response = client.post(
        "/job-search/compatibility-detail",
        headers=headers,
        json=_sample_listing_payload(),
    )

    assert response.status_code == 503

    app.dependency_overrides.pop(get_compatibility_detail_analyzer, None)


def test_compatibility_detail_rate_limited_after_max_per_hour(client):
    app.dependency_overrides[get_compatibility_detail_analyzer] = lambda: (
        FakeCompatibilityDetailAnalyzer()
    )
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    _onboard_with_cv(client, headers)

    for _ in range(MAX_COMPATIBILITY_DETAILS_PER_HOUR):
        response = client.post(
            "/job-search/compatibility-detail",
            headers=headers,
            json=_sample_listing_payload(),
        )
        assert response.status_code == 200

    response = client.post(
        "/job-search/compatibility-detail",
        headers=headers,
        json=_sample_listing_payload(),
    )
    assert response.status_code == 429

    app.dependency_overrides.pop(get_compatibility_detail_analyzer, None)
