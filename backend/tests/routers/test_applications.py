from app.llm_analyzer.analyzer import SemanticReport
from app.llm_analyzer.dependencies import get_semantic_analyzer
from app.main import app
from app.rate_limit.limiter import MAX_DIAGNOSTICS_PER_HOUR


class FakeAnalyzer:
    def analyze(self, cv_text, offer_text):
        return SemanticReport(
            score=70, missing_keywords=["Docker"], recommendations=["Add Docker"]
        )


def _register_and_login(client, email: str = "jane@example.com") -> str:
    client.post("/auth/register", json={"email": email, "password": "s3cret!1"})
    login = client.post("/auth/login", data={"username": email, "password": "s3cret!1"})
    return login.json()["access_token"]


def _setup_profile(client, headers: dict) -> None:
    client.put(
        "/profile",
        headers=headers,
        json={
            "full_name": "Jane Doe",
            "phone": "0600000000",
            "work_authorization": "FR/UE",
        },
    )
    import io

    from docx import Document

    document = Document()
    document.add_paragraph("Expérience professionnelle")
    document.add_paragraph("Développeuse Full Stack chez Acme, 2020-2022")
    document.add_paragraph("Formation")
    document.add_paragraph("Master Informatique")
    document.add_paragraph("Compétences")
    document.add_paragraph("Python, Docker")
    buffer = io.BytesIO()
    document.save(buffer)
    client.post(
        "/profile/cv",
        headers=headers,
        files={
            "cv_file": (
                "cv.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )


def test_create_list_and_get_application(client):
    app.dependency_overrides[get_semantic_analyzer] = lambda: FakeAnalyzer()
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    _setup_profile(client, headers)

    create = client.post(
        "/applications",
        headers=headers,
        json={
            "offer_url": "https://example.com/job/1",
            "offer_text": "Nous recherchons un développeur Python avec Docker.",
            "source": "manual",
            "company_name": "Acme",
            "job_title": "Développeur Python",
        },
    )
    assert create.status_code == 201
    body = create.json()
    assert body["status"] == "en_cours"
    assert body["diagnostic"]["missing_keywords"] == ["Docker"]
    application_id = body["id"]

    listing = client.get("/applications", headers=headers)
    assert listing.status_code == 200
    assert len(listing.json()) == 1

    detail = client.get(f"/applications/{application_id}", headers=headers)
    assert detail.status_code == 200
    assert detail.json()["company_name"] == "Acme"


def test_create_application_without_profile_cv_returns_422(client):
    app.dependency_overrides[get_semantic_analyzer] = lambda: FakeAnalyzer()
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}

    response = client.post(
        "/applications",
        headers=headers,
        json={
            "offer_url": "https://example.com/job/1",
            "offer_text": "Offre.",
            "source": "manual",
            "company_name": "Acme",
            "job_title": "Dev",
        },
    )
    assert response.status_code == 422


def test_create_duplicate_application_returns_409(client):
    app.dependency_overrides[get_semantic_analyzer] = lambda: FakeAnalyzer()
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    _setup_profile(client, headers)

    payload = {
        "offer_url": "https://example.com/job/1",
        "offer_text": "Offre.",
        "source": "manual",
        "company_name": "Acme",
        "job_title": "Dev",
    }
    first = client.post("/applications", headers=headers, json=payload)
    assert first.status_code == 201

    second = client.post("/applications", headers=headers, json=payload)
    assert second.status_code == 409


def test_get_application_not_owned_returns_404(client):
    app.dependency_overrides[get_semantic_analyzer] = lambda: FakeAnalyzer()
    owner_token = _register_and_login(client, "owner@example.com")
    owner_headers = {"Authorization": f"Bearer {owner_token}"}
    _setup_profile(client, owner_headers)
    created = client.post(
        "/applications",
        headers=owner_headers,
        json={
            "offer_url": "https://example.com/job/1",
            "offer_text": "Offre.",
            "source": "manual",
            "company_name": "Acme",
            "job_title": "Dev",
        },
    )
    application_id = created.json()["id"]

    attacker_token = _register_and_login(client, "attacker@example.com")
    attacker_headers = {"Authorization": f"Bearer {attacker_token}"}
    response = client.get(f"/applications/{application_id}", headers=attacker_headers)
    assert response.status_code == 404


def test_create_application_rate_limited_after_max_per_hour(client):
    app.dependency_overrides[get_semantic_analyzer] = lambda: FakeAnalyzer()
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    _setup_profile(client, headers)

    for i in range(MAX_DIAGNOSTICS_PER_HOUR):
        response = client.post(
            "/applications",
            headers=headers,
            json={
                "offer_url": f"https://example.com/job/{i}",
                "offer_text": "Offre.",
                "source": "manual",
                "company_name": "Acme",
                "job_title": "Dev",
            },
        )
        assert response.status_code == 201

    response = client.post(
        "/applications",
        headers=headers,
        json={
            "offer_url": "https://example.com/job/last",
            "offer_text": "Offre.",
            "source": "manual",
            "company_name": "Acme",
            "job_title": "Dev",
        },
    )
    assert response.status_code == 429
