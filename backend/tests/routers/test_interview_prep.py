from app.interview_prep.analyzer import InterviewPrepError
from app.interview_prep.dependencies import get_interview_prep_analyzer
from app.interview_prep.schemas import (
    CoachingChecklist,
    CompanyFacts,
    InterviewPrepDossierContent,
)
from app.llm_analyzer.analyzer import SemanticReport
from app.llm_analyzer.dependencies import get_semantic_analyzer
from app.main import app
from app.models.company_research_cache import CompanyResearchCache
from app.models.interview_prep_request_log import InterviewPrepRequestLog
from app.rate_limit.limiter import MAX_INTERVIEW_PREP_PER_HOUR

_OFFER_URL = "ftp://example.com/job/1"


class FakeInterviewPrepAnalyzer:
    def __init__(self, fail_draft: bool = False):
        self.research_calls = 0
        self.fail_draft = fail_draft

    def research_company(self, company_name, job_title):
        self.research_calls += 1
        return "Synthèse fictive.", [
            {"title": "Article", "url": "https://example.com/a"}
        ]

    def draft_dossier(
        self,
        cv_text,
        offer_text,
        missing_keywords,
        recommendations,
        persona,
        extra_context,
        company_research,
    ):
        if self.fail_draft:
            raise InterviewPrepError("boom")
        confidence = (
            "verified_web_search"
            if company_research
            else "general_knowledge_unverified"
        )
        return InterviewPrepDossierContent(
            narrative_angle="Un profil polyvalent.",
            company_facts=CompanyFacts(confidence=confidence),
            recent_news=[],
            probable_questions=[],
            practical_exercises=[],
            coaching_checklist=CoachingChecklist(before=[], during=[], after=[]),
        )


def _register_and_login(client, email: str = "jane@example.com") -> str:
    client.post("/auth/register", json={"email": email, "password": "s3cret!1"})
    login = client.post("/auth/login", data={"username": email, "password": "s3cret!1"})
    return login.json()["access_token"]


def _saved_job_payload(**overrides) -> dict:
    payload = {
        "offer_url": _OFFER_URL,
        "title": "Développeur Python",
        "company": "Acme",
        "location": "Paris",
        "snippet": "Poste de développeur.",
        "source": "adzuna",
        "ats_type": None,
        "salary": None,
    }
    payload.update(overrides)
    return payload


def _create_saved_job_with_diagnostic(client, headers) -> int:
    saved_id = client.post(
        "/saved-jobs", headers=headers, json=_saved_job_payload()
    ).json()["id"]

    import io

    from docx import Document

    class FakeSemanticAnalyzer:
        def analyze(self, cv_text: str, offer_text: str) -> SemanticReport:
            return SemanticReport(
                score=60, missing_keywords=["Docker"], recommendations=["Add Docker"]
            )

    document = Document()
    document.add_paragraph("Expérience professionnelle")
    document.add_paragraph("Développeur")
    document.add_paragraph("Formation")
    document.add_paragraph("Master")
    document.add_paragraph("Compétences")
    document.add_paragraph("Python")
    buffer = io.BytesIO()
    document.save(buffer)

    app.dependency_overrides[get_semantic_analyzer] = lambda: FakeSemanticAnalyzer()
    client.post(
        "/diagnostics",
        headers=headers,
        files={
            "cv_file": (
                "cv.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={
            "offer_text": "We need a Python developer.",
            "saved_job_id": str(saved_id),
        },
    )
    app.dependency_overrides.pop(get_semantic_analyzer, None)
    return saved_id


def _start_and_wait(client, headers, saved_job_id, **body):
    launch = client.post(
        f"/saved-jobs/{saved_job_id}/interview-prep", headers=headers, json=body
    )
    assert launch.status_code == 202
    job_id = launch.json()["job_id"]
    job = client.get(f"/generation-jobs/{job_id}", headers=headers)
    assert job.status_code == 200
    return job.json()


def test_interview_prep_requires_existing_diagnostic(client):
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    saved_id = client.post(
        "/saved-jobs", headers=headers, json=_saved_job_payload()
    ).json()["id"]

    response = client.post(
        f"/saved-jobs/{saved_id}/interview-prep",
        headers=headers,
        json={"persona": "coach bienveillant", "use_web_search": False},
    )
    assert response.status_code == 422


def test_interview_prep_without_web_search_completes(client):
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    saved_id = _create_saved_job_with_diagnostic(client, headers)

    fake = FakeInterviewPrepAnalyzer()
    app.dependency_overrides[get_interview_prep_analyzer] = lambda: fake

    job = _start_and_wait(
        client, headers, saved_id, persona="coach direct", use_web_search=False
    )
    assert job["status"] == "done"
    assert job["result"]["web_search_used"] is False
    assert fake.research_calls == 0

    detail = client.get(f"/saved-jobs/{saved_id}/interview-prep", headers=headers)
    assert detail.status_code == 200
    assert detail.json()["web_search_used"] is False

    app.dependency_overrides.pop(get_interview_prep_analyzer, None)


def test_interview_prep_with_web_search_caches_company_research(client, db_session):
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    saved_id = _create_saved_job_with_diagnostic(client, headers)

    fake = FakeInterviewPrepAnalyzer()
    app.dependency_overrides[get_interview_prep_analyzer] = lambda: fake

    job = _start_and_wait(
        client, headers, saved_id, persona="coach direct", use_web_search=True
    )
    assert job["status"] == "done"
    assert job["result"]["web_search_used"] is True
    assert fake.research_calls == 1
    assert db_session.query(CompanyResearchCache).count() == 1

    # Regenerating for the same (cached) company must not re-trigger research.
    _start_and_wait(
        client, headers, saved_id, persona="coach direct", use_web_search=True
    )
    assert fake.research_calls == 1

    app.dependency_overrides.pop(get_interview_prep_analyzer, None)


def test_get_interview_prep_404s_before_generation_and_for_other_users(client):
    token_a = _register_and_login(client, "a@example.com")
    token_b = _register_and_login(client, "b@example.com")
    headers_a = {"Authorization": f"Bearer {token_a}"}
    headers_b = {"Authorization": f"Bearer {token_b}"}
    saved_id = _create_saved_job_with_diagnostic(client, headers_a)

    before = client.get(f"/saved-jobs/{saved_id}/interview-prep", headers=headers_a)
    assert before.status_code == 404

    fake = FakeInterviewPrepAnalyzer()
    app.dependency_overrides[get_interview_prep_analyzer] = lambda: fake
    _start_and_wait(
        client, headers_a, saved_id, persona="coach direct", use_web_search=False
    )
    app.dependency_overrides.pop(get_interview_prep_analyzer, None)

    other_user = client.get(f"/saved-jobs/{saved_id}/interview-prep", headers=headers_b)
    assert other_user.status_code == 404


def test_interview_prep_rate_limit_returns_429_after_cap(client):
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    saved_id = _create_saved_job_with_diagnostic(client, headers)

    app.dependency_overrides[get_interview_prep_analyzer] = lambda: (
        FakeInterviewPrepAnalyzer()
    )
    for _ in range(MAX_INTERVIEW_PREP_PER_HOUR):
        response = client.post(
            f"/saved-jobs/{saved_id}/interview-prep",
            headers=headers,
            json={"persona": "coach direct", "use_web_search": False},
        )
        assert response.status_code == 202

    over_cap = client.post(
        f"/saved-jobs/{saved_id}/interview-prep",
        headers=headers,
        json={"persona": "coach direct", "use_web_search": False},
    )
    assert over_cap.status_code == 429

    app.dependency_overrides.pop(get_interview_prep_analyzer, None)


def test_interview_prep_job_ends_in_error_and_does_not_consume_quota(
    client, db_session
):
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    saved_id = _create_saved_job_with_diagnostic(client, headers)

    app.dependency_overrides[get_interview_prep_analyzer] = lambda: (
        FakeInterviewPrepAnalyzer(fail_draft=True)
    )
    job = _start_and_wait(
        client, headers, saved_id, persona="coach direct", use_web_search=False
    )
    assert job["status"] == "error"
    assert db_session.query(InterviewPrepRequestLog).count() == 0

    app.dependency_overrides.pop(get_interview_prep_analyzer, None)
